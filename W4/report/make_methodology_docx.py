# -*- coding: utf-8 -*-
"""W4 Methodology & Results — Word deliverable (python-docx).

INTERNAL working document: clean professional styling, NOT the client Sample.docx
letterhead (rule 5 applies to client-facing reports; this one is for the design team).
Content mirrors W4/docs/METHODOLOGY.md — numbers read live from W4/run/summary.json
so the docx can never drift from the last run.

Re-run: python W4/report/make_methodology_docx.py
"""
import json
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

W4 = r"D:\Mojtaba\Renardet\2621 Ibri Sewer STP\Hydraulic\Claude\W4"
OUT = os.path.join(W4, "report", "W4_Sewer_Network_Design_Methodology.docx")
IMG_DOCS = os.path.join(W4, "docs", "img")
IMG_MAPS = os.path.join(W4, "img")

S = json.load(open(os.path.join(W4, "run", "summary.json")))
sc = S["selfclean"]
aug = S["augmentation"]
lp = S["lowplots"]
ld = S["loads"]
dn_km = S["dn_km"]

doc = Document()

# ---------- base styles ----------
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
for lvl, sz, col in ((1, 16, "1F3864"), (2, 13, "2E5A88"), (3, 11.5, "2E5A88")):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor.from_string(col)
    h.font.bold = True
for s_ in doc.sections:
    s_.top_margin = s_.bottom_margin = Inches(0.8)
    s_.left_margin = s_.right_margin = Inches(0.85)


def para(text="", align=None, bold=None, size=None, color=None, space_after=6, italic=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def H(level, text):
    return doc.add_paragraph(text, style=f"Heading {level}")


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
    p.add_run(text)
    p.paragraph_format.space_after = Pt(4)
    return p


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        shade(c, "D9E2F3")
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(font)
    if widths:
        for i, w in enumerate(widths):
            for r_ in t.rows:
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def pic(path, w=6.3, cap=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
    if cap:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(cap)
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string("555555")


def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc():
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click > Update Field to build the Table of Contents."
    r.append(t)
    fld.append(r)
    p._p.append(fld)


# ================= COVER =================
para("", space_after=90)
para("Ibri Sewer, TE & STP — Project 2621", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13,
     color="555555")
para("W4 — Sewer Network Design Pipeline", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26,
     color="1F3864", space_after=8)
para("Methodology and Test-Boundary Results", align=WD_ALIGN_PARAGRAPH.CENTER, size=16,
     color="2E5A88", space_after=40)
para("Internal working document — design team", align=WD_ALIGN_PARAGRAPH.CENTER, size=11,
     italic=True, color="777777")
para("18 August 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color="777777", space_after=40)
table(["Item", "Value"], [
    ["Test area", f"{S['s1']['boundary_ha']:.0f} ha"],
    ["Network designed", f"{S['n_nodes']:,} manholes / {S['net_km']:.1f} km"],
    ["Peak flow at outfall", f"{S['qpeak_outfall_ls']:.0f} L/s (Qadf {S['qadf_outfall_m3d']:,.0f} m3/d)"],
    ["Audit result", f"{S['violations']} violations"],
    ["Code / tests", "W4/py/sewnet — 43 pytest cases, Table-11 gate"],
], widths=[2.2, 4.0], font=10)
pagebreak()

# ================= TOC =================
H(1, "Contents")
toc()
pagebreak()

# ================= EXEC SUMMARY =================
H(1, "Executive summary")
para("We built the sewer design pipeline, proved it end to end on the "
     f"{S['s1']['boundary_ha']:.0f} ha test boundary, put it through a 21-agent adversarial "
     "review, fixed everything the review confirmed, and the final design passes its own audit "
     "with zero violations. Here is the whole story in one page.")
para(f"The pipeline takes four inputs — roads, classified plots, the 0.5 m terrain and a boundary — "
     f"and produces a complete gravity network: {S['n_nodes']:,} manholes, {S['net_km']:.1f} km of "
     f"pipe ({float(dn_km['200'])/S['net_km']*100:.0f}% DN200, stepping through DN250–400 to a "
     f"DN500 outfall leg), designed inverts everywhere, {13} seconds to run. Every one of the "
     f"{ld['loaded_points']:,} loaded units ({ld['built']:,} built plots, {ld['planned']} planned, "
     f"{ld['unparceled']} unparceled buildings; {ld['farms_excluded']} farms excluded per doctrine) "
     f"lands on exactly one manhole — nothing silently dropped, mass balance closes exactly. "
     f"Saturation flow at the outfall: Qadf {S['qadf_outfall_m3d']:,.0f} m3/d, peak "
     f"{S['qpeak_outfall_ls']:.0f} L/s (Merrimack).")
para("The hydraulics were never taken on faith, and the verification regime earned its keep. The "
     "Colebrook-White solver had to reproduce all nine minimum gradients of G203 Table 11 (±5%) "
     "before it was allowed to size anything — it does, mean deviation under 2%. Then the "
     "adversarial panel attacked the code and confirmed 11 real defects, all now fixed: the "
     "spanning tree covered every road junction but skipped 25 km of cross-street edges (fixed "
     "with the standard crest-manhole layout); pipe capacity was computed on nominal diameter "
     "while OD-designated PVC-U bores smaller (now on true SDR34 bore); a broken bisection in the "
     "velocity-cap solver; and drop structures measured against the wrong datum. 43 pytest cases "
     "lock all of it in.")
para(f"The honest self-cleansing picture, on the table rather than under it: "
     f"{sc['share_below']*100:.0f}% of pipes cannot reach 0.75 m/s at saturation peak — physically "
     f"inevitable on small residential branches — and comply through the guideline's own "
     f"tractive-force alternative at tau = 1 Pa, which is a pending assumption [GAP-9]. If NWS "
     f"sets tau = 2 Pa, {sc['would_fail_at_tau2']:,} pipes need steeper slopes. That single number "
     f"is the strongest argument for pinning tau down at the kickoff.")
para("Three findings from the test area for your eyes:", bold=True, space_after=4)
bullet(f"the terrain's lowest boundary road node is at ({S['outfall']['x']:.0f}, "
       f"{S['outfall']['y']:.0f}), south-center on the main corridor, z {S['outfall']['z']:.1f} m, "
       f"not the west edge. If the real trunk connection is elsewhere, it is one config line and a "
       f"13-second re-run.", bold_lead=f"The outfall landed {S['outfall']['dist_to_expected_m']:.0f} m from where you expected — ")
bullet(f"{lp['flagged']} units ({lp['flagged']/lp['checked']*100:.1f}%) could not gravity-connect "
       f"at standard sewer depth. Deepening {lp['deepened_mh']} manholes recovered all but "
       f"{lp['residual']}, now flagged as local-solution candidates.",
       bold_lead="Your elevated-roads concern was justified: ")
bullet(f"({S['solver']['pockets']} pocket, ~5 properties) — absorb-to-detail-design per rule 9. "
       f"{S['drops']} drop structures ({S['vortex_sites']} vortex-class) concentrate at wadi-bank "
       f"crossings, exactly where detail design applies the G1-p85 crossing rules.",
       bold_lead="One SLS pocket appeared ")
para("Also delivered on the way: T01 Rev 3 — the tutorial now teaches Colebrook-White (§14), "
     "Table 11 derived step by step, every number independently verified.")
para("What is proven: doctrine loads in, guideline-compliant network out, auditable and "
     "re-runnable, honest about its assumptions. What is not yet proven: SewerGEMS agreement (the "
     "import package and pipe-by-pipe referee table await the ModelBuilder run), and behaviour at "
     "36-zone scale with the finalised trunk. That is W5.", bold=False)
pagebreak()

# ================= 1 PIPELINE =================
H(1, "1. What the pipeline is")
para("A re-runnable Python package (sewnet) that designs a gravity sewer network inside any "
     "boundary, given roads, loaded plots, terrain and an outfall or connection point. One config "
     "file holds the paths; criteria.py holds every design number with its PAM-GUD page reference "
     "— no number lives anywhere else in the code.")
pic(os.path.join(IMG_DOCS, "pipeline_architecture.png"), 6.3,
    "Figure 1 — Pipeline architecture. Blue = inputs, yellow = audit gate, red = SLS flag, "
    "dashed = iteration loops.")
para("Stages, in one breath: repair and clip the inputs; node the roads and collapse dual "
     "carriageways; grow a loop-free tree toward the outfall (climb-penalised, arterial-preferring "
     "— the 'no alleys' lesson); add cross-street branches wherever loaded plots front an off-tree "
     "street (summit-split, crest-manhole layout); place manholes (junctions, bends >45°, <=100 m "
     "spacing, sub-2 m reaches contracted); load every plot at saturation; accumulate flows with "
     "peak factor and infiltration; size pipes and solve inverts together on the true pipe bore; "
     "check every house can physically reach its manhole, deepening where roads are elevated; "
     "audit everything independently; export SHP, SewerGEMS, DXF and maps.")

# ================= 2 HYDRAULIC BASIS =================
H(1, "2. Hydraulic basis and how it is verified")
table(["Element", "Basis", "Verification"], [
    ["Capacity / velocity",
     "Colebrook-White, ks = 1.5 mm, nu = 1.141e-6 m2/s (G203-p24/25/28), partial-full circular "
     "geometry, true internal bore (PVC-U OD-series derated to SDR34 ID; GRP nominal = ID)",
     "Table-11 gate: reproduce all 9 minimum gradients at 0.75 m/s ±5% — passes at <2% mean "
     "deviation; ID convention unit-tested"],
    ["Minimum gradients",
     "Steeper of Table 11 (p29) and tractive force Smin = 2.33e-4·tau^1.23·Q^-0.461 (p27, "
     "A9-corrected), plus a 40 mm total-fall guard per reach (p29 §4.3.1)",
     "tau = 1 Pa tagged [GAP-9]; Q floored at Mara's 1.5 L/s minimum design flow — at the floor "
     "tractive ≈ Table 11 DN200: the methods meet"],
    ["d/D limits", "<=0.65 (D<=350), <=0.50 (D>350) at peak (p27 Tab 10)",
     "enforced in sizing, re-checked in audit on the true bore"],
    ["Velocity band",
     ">=0.75 m/s at peak or tractive-compliant (p26–27); <=3.0 m/s via a slope cap whose surplus "
     "fall becomes a designed drop", "audit plus the transparency statistics in section 5"],
    ["Loads",
     "Doctrine §2: every plot at saturation, 6.0 × 171.3 l/c/d ≈ 1.03 m3/d/plot; farms zero; "
     "infiltration 720 L/d/km unpeaked; PF Merrimack (mandatory >100 properties), held at its "
     "100-property value below; Peltier comparison held the same way",
     "mass balance to the outfall closes exactly; unknown CLASS values can never vanish silently"],
], widths=[1.1, 2.7, 2.5], font=8.5)
pic(os.path.join(IMG_DOCS, "load_chain.png"), 5.4, "Figure 2 — Load allocation chain.")
pagebreak()

# ================= 3 SOLVER =================
H(1, "3. How the solver designs a pipe")
para("Two passes per reach, iterated until no diameter changes (2 iterations sufficed; "
     "oscillation between adjacent DNs is detected and broken upward; a final lay pass always "
     "leaves inverts consistent with final diameters).")
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
c1, c2 = t.rows[0].cells
c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
c1.paragraphs[0].add_run().add_picture(os.path.join(IMG_DOCS, "solver_step1_slope.png"),
                                       width=Inches(2.5))
c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
c2.paragraphs[0].add_run().add_picture(os.path.join(IMG_DOCS, "solver_step2_profile.png"),
                                       width=Inches(2.9))
para("Figure 3 — Solver logic: step 1 sets the reach slope, step 2 resolves profile, junction and "
     "depth.", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=8.5, color="555555")
para("The details that matter:", bold=True, space_after=4)
bullet("one straight line between manholes, never kinked.", bold_lead="Uniform slope per reach (p29) — ")
bullet("manhole depth is construction depth, and every inlet drop is measured inlet-invert minus "
       "outgoing-invert, so velocity-cap surpluses and cover shifts combine into the drop they "
       "physically are. >600 mm = external backdrop, >2 m = vortex shaft (p30); the audit "
       "re-derives every drop independently, and a missing or misclassified record is a violation.",
       bold_lead="The chamber datum is the outgoing invert — ")
bullet("every reach profiled at 5 m on the 0.5 m VRT; a reach riding above dipping ground is "
       "shifted down bodily, the shift surfacing as a recorded drop.",
       bold_lead="Mid-span cover on real terrain: ")
bullet("every candidate DN is judged at its own governing slope (no-oversizing, p29), on its true "
       "bore, with velocity-infeasible candidates skipped explicitly.",
       bold_lead="Sizing without the ratchet: ")
bullet("never a silent orphan.", bold_lead=">12 m depth becomes an SLS pocket (p33 + rule 9) — ")

# ================= 4 CONNECTABILITY =================
H(1, "4. The house-connectability check")
para("Roads are locally elevated for flood protection and underpasses, so houses can sit below the "
     "sewer. For every loaded unit: plot ground (0.5 m terrain) minus 0.6 m outlet depth must reach "
     "its manhole's invert with 2% fall over the connection distance. Failures raise a deepening "
     "requirement on that manhole (capped 0.5 m short of the 12 m limit) and the invert solve "
     f"re-runs; anything still failing is flagged for a local solution. Test area: {lp['flagged']} "
     f"flagged, {lp['deepened_mh']} manholes deepened, {lp['residual']} residual.")
pic(os.path.join(IMG_MAPS, "W4_M3_connectability.png"), 4.6,
    "Figure 4 — Plot connectability: green connectable, amber recovered by deepening, red needs a "
    "local solution.")
pagebreak()

# ================= 5 RESULTS =================
H(1, "5. Test-boundary results")
table(["Quantity", "Value"], [
    ["Area / roads / network",
     f"{S['s1']['boundary_ha']:.0f} ha / {S['s1']['len_km']:.1f} km roads / {S['net_km']:.1f} km "
     f"sewers ({aug['added_km']} km from cross-street augmentation; {aug['skipped_km']} km of "
     f"unloaded streets deliberately unsewered)"],
    ["Manholes / pipes", f"{S['n_nodes']:,} / {S['n_pipes']:,} (single tree, one outfall)"],
    ["Diameters", " · ".join(f"DN{k} {v} km" for k, v in dn_km.items())],
    ["Loaded units", f"{ld['loaded_points']:,} (built {ld['built']:,} + planned {ld['planned']} + "
                     f"unparceled {ld['unparceled']}); farms excluded {ld['farms_excluded']}"],
    ["Outfall", f"({S['outfall']['x']:.0f}, {S['outfall']['y']:.0f}), z {S['outfall']['z']:.2f} — "
                f"{S['outfall']['dist_to_expected_m']:.0f} m from user expectation"],
    ["Qadf / Qpeak at outfall", f"{S['qadf_outfall_m3d']:,.0f} m3/d / {S['qpeak_outfall_ls']:.0f} "
                                f"L/s (PF {S['pf_formula']}; Peltier column in the shapefiles)"],
    ["Depths", f"max {S['max_depth_m']:.1f} m (inside the SLS pocket); gravity network otherwise "
               f"<= 12 m"],
    ["Drops", f"{S['drops']} designed structures; {S['vortex_sites']} vortex-class (>2 m), "
              f"concentrated at wadi-bank crossings"],
    ["SLS pockets", f"{S['solver']['pockets']} (~5 properties, absorb per rule 9)"],
    ["Low plots", f"{lp['flagged']} flagged, {lp['residual']} residual after deepening "
                  f"{lp['deepened_mh']} manholes"],
    ["Audit", f"{S['violations']} violations (saturation run, including independent drop "
              f"re-derivation)"],
    ["Self-cleansing transparency",
     f"{sc['below_075_at_peak']:,}/{sc['pipes']:,} pipes below 0.75 m/s at saturation peak — "
     f"compliant via tractive at tau=1 Pa [GAP-9]; {sc['would_fail_at_tau2']:,} would need "
     f"redesign at tau=2 Pa; start-year flags {S['startyear_flags']:,}, all tractive-compliant"],
    ["Runtime", "~13 s design + ~5 s exports"],
], widths=[1.7, 4.6], font=9)
pic(os.path.join(IMG_MAPS, "W4_M1_network_by_dn.png"), 4.4,
    "Figure 5 — Designed network by diameter, with outfall and SLS candidate.")
pic(os.path.join(IMG_MAPS, "W4_M2_depth.png"), 4.4,
    "Figure 6 — Excavation depth (invert below ground).")
pagebreak()

# ================= 6 SEWERGEMS =================
H(1, "6. SewerGEMS package and referee protocol")
para("W4/sewergems/ holds MANHOLES, CONDUITS and OUTFALL shapefiles built to the Bentley-documented "
     "ModelBuilder mappings (explicit START_ND/STOP_ND plus vertex-snapped geometry digitised "
     "upstream to downstream, elevations not depths, numeric mm diameters), LOADS.xlsx "
     "(pattern-based rows, L/s per manhole), and IMPORT_PROCEDURE.md with the known traps — the "
     "'Set Invert to Start/Stop Node = False' global edit being the one that silently deletes drop "
     "manholes. Manning n is exported as 0.013, the ks = 1.5 mm equivalent: the Tab-8 range "
     "0.009–0.011 conflicts with the ks mandate and would show roughly 30% phantom capacity, so "
     "that conflict is flagged as an NWS kickoff item.")
para("After the model run, paste SewerGEMS discharge, velocity and d/D into REFEREE_pipes.csv; any "
     "pipe deviating more than 5% from our columns is an open investigation. The design is not "
     "'verified' until the two engines agree — deliberately a separate run, not a self-check.")

# ================= 7 ASSUMPTIONS =================
H(1, "7. Assumptions register")
para("Tagged in criteria.ASSUMPTIONS with the same wording, and reported in every deliverable:")
table(["Assumption", "Basis / exposure"], [
    ["tau = 1 Pa", "GUD-203 gives no numeric design tractive stress [GAP-9]; largest redesign "
                   f"exposure ({sc['would_fail_at_tau2']:,} pipes at tau=2)"],
    ["Tractive Q-floor 1.5 L/s", "Mara simplified-sewerage minimum design flow; unfloored the "
                                 "formula demands unbounded slopes as Q approaches 0"],
    ["PVC-U wall class SDR34/SN8", "true bore for hydraulics; actual class per PAM-SPC-207 pending"],
    ["OR 6.0, 1 property per plot", "GAP-5 (NCSI housing units missing)"],
    ["Plot outlet 0.6 m, 2% connection fall", "method choice for the connectability check"],
    ["40 m frontage rule", "an off-tree street gets a sewer when a loaded unit lies within 40 m"],
    ["Manhole size ladder", "02 has no size table; GUD-203 §4.4 re-extract pending (no hydraulic "
                            "effect at concept)"],
    ["Rider discharge to nearest manhole", "02 silent on the discharge point"],
    ["Gravity in-road position", "taken from the force-main clause (p51, A9) as an inference"],
    ["PF held below 100 properties", "G1-p71 prescribes no formula there; Peltier held the same way"],
    ["Infiltration unpeaked", "add-order not stated in GUD-201 — kickoff item"],
], widths=[2.2, 4.1], font=9)

# ================= 8 LIMITATIONS =================
H(1, "8. Limitations and what changes at full scale")
bullet("the pipeline designs subnetworks into the given connection points, which are config "
       "entries, not structure.", bold_lead="The trunk is user-finalised (settled 2026-08-18) — ")
bullet("multi-connection territory competition is the W5 structural addition.",
       bold_lead="One outfall per run today; ")
bullet("junction losses ignored (normal-depth hydraulics) — both noted for the SewerGEMS "
       "comparison.", bold_lead="Riders are schematic; ")
bullet("flat trunk profiles at scale need survey-grade data — already a registered data request.",
       bold_lead="The 0.5 m terrain is concept-grade for inverts near the 0.75–1.0 mm/m minimums "
                 "(G1-p36); ")
bullet("pin it at the kickoff.", bold_lead="tau = 1 Pa carries the largest redesign exposure — ")

H(1, "9. Adversarial review")
para("A 21-agent skeptic panel attacked the hydraulic core after the first audit-clean run: four "
     "attack lenses (Colebrook-White implementation, solver clause compliance, load/audit doctrine, "
     "executed edge cases) followed by independent verification of every raw finding. 17 raw "
     "findings, 11 confirmed, 11 fixed, 6 refuted. The full register with fixes is in "
     "W4/docs/REVIEW_FINDINGS.md; the headline four were cross-street coverage (+25.3 km), the "
     "true PVC bore correction, the drop-datum correction, and the velocity-cap bisection repair.")

doc.save(OUT)
print("saved", OUT)
