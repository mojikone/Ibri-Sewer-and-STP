# -*- coding: utf-8 -*-
"""Tutorial T01 Rev 4 — Sewage Flow & Load Calculation, styled on Data/sample report/Sample.docx.
Rev 1 addresses the review comments of T01_..._commented.docx (2026-08-14).
Rev 4 (2026-08-19): properties per plot are now COUNTED from electricity accounts and the
occupancy rate is set to 5 — Route B is no longer an assumed product.
Rev 3 (2026-08-18): added the Colebrook-White pipe-hydraulics chapter (Section 14).
Refs: G203-p## = PAM-GUD-203, G201-p## = PAM-GUD-201, G202-p## = PAM-GUD-202,
R0 = Ibri Inception Report R0 demand workbook (Aug 2026)."""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

SAMPLE = r"D:\Mojtaba\Renardet\2621 Ibri Sewer STP\Data\sample report\Sample.docx"
HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "T01_Sewage_Flow_and_Load_Calculation.docx")
IMG = os.path.join(HERE, "img")
os.makedirs(IMG, exist_ok=True)
PROJ = "Consultancy Services for Design and Supervision for STP, Sewer & TE Networks Systems in Ibri"
SUB = "Tutorial T01 — Sewage Flow and Pollution Load Calculation"

BLUE = "#0072B2"; ORANGE = "#E69F00"; INK = "#333333"; GRID = "#DDDDDD"

# ================================================================= CHARTS
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np

def style_ax(ax):
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    for s in ("left", "bottom"):
        ax.spines[s].set_color("#999999")
    ax.grid(True, color=GRID, lw=0.7, zorder=0)
    ax.tick_params(colors="#555555", labelsize=9)

CHAIN = os.path.join(IMG, "t01_chain.png")
def make_chain():
    steps = [
        ("1  Population", "NCSI census route / plot route / electricity pro-rata"),
        ("2  Water demand", "domestic 164 l/c/d + non-domestic 22% + governmental 14%"),
        ("3  Wastewater return", "85% of domestic & tanker / 54% of non-domestic & governmental"),
        ("4  Additions", "infiltration + tankered sewage"),
        ("5  Average flow Qadf", "the annual-average flow (AAF)"),
        ("6  Peak flows", "Peltier / Merrimack peak factors, PF <= 5"),
        ("7  Organic loads", "60 g BOD5, 80 g TSS per capita per day"),
        ("8  STP sizing", "biology on AAF + loads; hydraulics on PHF; +10% margin"),
        ("9  TSE output", "95% of STP inlet, feeds the TE network"),
    ]
    fig, ax = plt.subplots(figsize=(7.4, 8.6))
    ax.set_xlim(0, 10); ax.set_ylim(0, len(steps) * 1.15); ax.axis("off")
    ys = []
    for i, (t, s) in enumerate(steps):
        y = (len(steps) - 1 - i) * 1.15
        ax.add_patch(FancyBboxPatch((0.9, y + 0.08), 8.2, 0.86, boxstyle="round,pad=0.06",
                                    fc="#EAF1FB", ec="#2E5C9E", lw=1.4))
        ax.text(5.0, y + 0.63, t, ha="center", va="center", fontsize=12.5, fontweight="bold", color="#1F3A64")
        ax.text(5.0, y + 0.28, s, ha="center", va="center", fontsize=9.5, color="#444444")
        ys.append(y)
    for i in range(len(steps) - 1):
        ax.add_patch(FancyArrowPatch((5.0, ys[i] + 0.06), (5.0, ys[i + 1] + 0.96),
                                     arrowstyle="-|>", mutation_scale=16, color="#2E5C9E", lw=1.4))
    fig.savefig(CHAIN, dpi=200, bbox_inches="tight"); plt.close(fig)

DIURNAL = os.path.join(IMG, "t01_diurnal.png")
def make_diurnal():
    h = np.linspace(0, 24, 241)
    f = (1.0 + 0.55 * np.sin((h - 9.5) * np.pi / 12) + 0.28 * np.sin((h - 20.0) * np.pi / 6.0)
         - 0.18 * np.cos((h - 3) * np.pi / 12))
    f = np.clip(f, 0.25, None)
    f = f / f.mean()
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.plot(h, f, color=BLUE, lw=2.2, zorder=3)
    ax.axhline(1.0, color="#777777", lw=1.2, ls="--", zorder=2)
    pk = h[np.argmax(f)]
    ax.annotate("peak hour  (PHF = peak-hour flow)", xy=(pk, f.max()), xytext=(pk + 2.2, f.max() + 0.06),
                fontsize=9, color=INK, arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.annotate("daily average = 1.0  (AAF when averaged over the year)", xy=(2.8, 1.0),
                xytext=(0.4, 1.62), fontsize=9, color=INK, arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.annotate("night minimum\n(mostly infiltration)", xy=(h[np.argmin(f)], f.min()),
                xytext=(6.5, 0.32), fontsize=9, color=INK, arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.set_xlim(0, 24); ax.set_xticks(range(0, 25, 4))
    ax.set_xlabel("hour of day", fontsize=9, color=INK)
    ax.set_ylabel("flow / daily average", fontsize=9, color=INK)
    style_ax(ax)
    fig.savefig(DIURNAL, dpi=200, bbox_inches="tight"); plt.close(fig)

ROUTE_CENSUS = os.path.join(IMG, "t01_route_census.png")
ROUTE_PLOTS = os.path.join(IMG, "t01_route_plots.png")
def route_chart(path, title, steps, color="#2E5C9E", fill="#EAF1FB"):
    fig, ax = plt.subplots(figsize=(7.2, 1.1 + 1.12 * len(steps)))
    ax.set_xlim(0, 10); ax.set_ylim(0, len(steps) * 1.12 + 0.5); ax.axis("off")
    ax.text(5, len(steps) * 1.12 + 0.25, title, ha="center", fontsize=12.5, fontweight="bold", color="#1F3A64")
    ys = []
    for i, (t, s) in enumerate(steps):
        y = (len(steps) - 1 - i) * 1.12
        ax.add_patch(FancyBboxPatch((0.7, y + 0.10), 8.6, 0.84, boxstyle="round,pad=0.05",
                                    fc=fill, ec=color, lw=1.3))
        ax.text(5.0, y + 0.62, t, ha="center", va="center", fontsize=10.5, fontweight="bold", color="#1F3A64")
        ax.text(5.0, y + 0.29, s, ha="center", va="center", fontsize=8.8, color="#444444")
        ys.append(y)
    for i in range(len(steps) - 1):
        ax.add_patch(FancyArrowPatch((5.0, ys[i] + 0.08), (5.0, ys[i + 1] + 0.95),
                                     arrowstyle="-|>", mutation_scale=14, color=color, lw=1.3))
    fig.savefig(path, dpi=200, bbox_inches="tight"); plt.close(fig)

POPCHART = os.path.join(IMG, "t01_population.png")
SETTLCHART = os.path.join(IMG, "t01_pop_settlements.png")
NCSI = {2021: 168409, 2022: 173418, 2023: 178477, 2024: 183564, 2025: 187962, 2026: 193116,
        2027: 198344, 2028: 203659, 2029: 209063, 2030: 213637, 2031: 219186, 2032: 224840, 2033: 227736}
import json
_ps_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "pop_series.json")
POP_SERIES = json.load(open(_ps_path)) if os.path.exists(_ps_path) else None
def _sv(name, year):
    return POP_SERIES[name][str(year)] if POP_SERIES else None
def make_pop():
    yrs = sorted(NCSI); pop = [NCSI[y] / 1000 for y in yrs]
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    ax.plot(yrs, pop, color=BLUE, lw=2.2, marker="o", ms=4.5, zorder=3)
    ax.annotate("183.6k (2024)", xy=(2024, 183.564), xytext=(2024.2, 172), fontsize=9, color=INK,
                arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.annotate("227.7k (2033)", xy=(2033, 227.736), xytext=(2030.4, 231), fontsize=9, color=INK)
    ax.text(2021.2, 224, "growth rate declines 3.0% -> 1.3%/yr", fontsize=9, color="#555555")
    ax.set_xlabel("year", fontsize=9, color=INK); ax.set_ylabel("population (thousands)", fontsize=9, color=INK)
    ax.set_xticks(yrs[::2])
    style_ax(ax)
    fig.savefig(POPCHART, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_pop_settlements():
    series = [("IBRI", BLUE), ("AD DARIZ", ORANGE), ("AL ARAQI", "#009E73"), ("TOTAL", "#444444")]
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    for name, col in series:
        yrs = sorted(int(y) for y in POP_SERIES[name])
        vals = [POP_SERIES[name][str(y)] / 1000 for y in yrs]
        ls = "--" if name == "TOTAL" else "-"
        ax.plot(yrs, vals, color=col, lw=2.0, ls=ls, zorder=3)
        lbl = "Project TOTAL (50 settlements)" if name == "TOTAL" else name
        dy = {"AD DARIZ": 9, "AL ARAQI": -9}.get(name, 0)
        ax.annotate(lbl, xy=(yrs[-1], vals[-1]), xytext=(4, dy), textcoords="offset points",
                    fontsize=8.6, color=col, fontweight="bold", va="center")
    ax.set_xlim(2023, 2126)
    ax.set_xticks(range(2030, 2101, 10))
    ax.set_xlabel("year", fontsize=9, color=INK)
    ax.set_ylabel("population (thousands)", fontsize=9, color=INK)
    style_ax(ax)
    fig.savefig(SETTLCHART, dpi=200, bbox_inches="tight"); plt.close(fig)

_extra_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data", "r0_extra.json")
EXTRA = json.load(open(_extra_path)) if os.path.exists(_extra_path) else None
DEMAND_BUILDUP = os.path.join(IMG, "t01_demand_buildup.png")
RETURN_CHART = os.path.join(IMG, "t01_return.png")
RAMP_CHART = os.path.join(IMG, "t01_ramp.png")
TANKER_CHART = os.path.join(IMG, "t01_tankers.png")
LPCD_CHART = os.path.join(IMG, "t01_lpcd.png")
WWG_CHART = os.path.join(IMG, "t01_wwg.png")

def make_demand_buildup():
    fig, ax = plt.subplots(figsize=(7.0, 1.9))
    segs = [("Domestic  164", 164.0, BLUE), ("Non-domestic  36.1", 36.1, ORANGE), ("Governmental  23.0", 23.0, "#009E73")]
    left = 0
    for lbl, v, col in segs:
        ax.barh([0], [v], left=left, color=col, height=0.5, zorder=3)
        ax.text(left + v / 2, 0.45, lbl, ha="center", fontsize=9, color=col, fontweight="bold")
        left += v
    ax.text(left + 3, 0, f"total  {left:.0f} l/c/d", va="center", fontsize=10, color=INK, fontweight="bold")
    ax.set_xlim(0, 265); ax.set_ylim(-0.6, 0.9); ax.set_yticks([])
    ax.set_xlabel("litres per capita per day", fontsize=9, color=INK)
    style_ax(ax); ax.grid(False); ax.spines["left"].set_visible(False)
    fig.savefig(DEMAND_BUILDUP, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_return_chart():
    cats = ["Domestic", "Non-dom + gov", "TOTAL"]
    supplied = [164.0, 59.1, 223.1]
    returned = [139.4, 31.9, 171.3]
    x = np.arange(3); w = 0.34
    fig, ax = plt.subplots(figsize=(7.0, 3.0))
    ax.bar(x - w / 2, supplied, w, color="#9CC7E4", zorder=3, label="water supplied")
    ax.bar(x + w / 2, returned, w, color=BLUE, zorder=3, label="returned as sewage")
    for i in range(3):
        ax.text(x[i] - w / 2, supplied[i] + 4, f"{supplied[i]:.0f}", ha="center", fontsize=9, color="#6A93B5")
        ax.text(x[i] + w / 2, returned[i] + 4, f"{returned[i]:.0f}", ha="center", fontsize=9, color=BLUE, fontweight="bold")
    ax.set_xticks(x); ax.set_xticklabels(cats, fontsize=9.5)
    ax.set_ylabel("l/c/d", fontsize=9, color=INK); ax.set_ylim(0, 260)
    ax.legend(fontsize=8.5, frameon=False, loc="upper left")
    style_ax(ax)
    fig.savefig(RETURN_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_ramp_chart():
    ramp = {int(y): v for y, v in EXTRA["ramp"].items()}
    yrs = sorted(ramp); vals = [ramp[y] * 100 for y in yrs]
    fig, ax = plt.subplots(figsize=(7.0, 2.9))
    ax.plot(yrs, vals, color=BLUE, lw=2.2, zorder=3)
    ax.axhline(100, color="#999999", lw=1.0, ls=":")
    ax.annotate("46.7% (2021)", xy=(2021, 46.7), xytext=(2024, 36), fontsize=9, color=INK,
                arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.annotate(f"100% reached {EXTRA['ramp_100']}", xy=(EXTRA["ramp_100"], 100), xytext=(EXTRA["ramp_100"] - 22, 82),
                fontsize=9, color=INK, arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.set_xlim(2020, 2101); ax.set_ylim(30, 112)
    ax.set_xlabel("year", fontsize=9, color=INK); ax.set_ylabel("connected population (%)", fontsize=9, color=INK)
    style_ax(ax)
    fig.savefig(RAMP_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_tanker_chart():
    top = EXTRA["tanker_top"][::-1]
    names = [t[0] for t in top]; vals = [t[1] for t in top]
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.barh(names, vals, color=BLUE, height=0.62, zorder=3)
    for i, v in enumerate(vals):
        ax.text(v + 3, i, f"{v:.0f}", va="center", fontsize=8.6, color=INK)
    ax.set_xlabel("tankered volume 2024 (m3/d, held constant in R0)", fontsize=9, color=INK)
    ax.tick_params(axis="y", labelsize=8.6)
    ax.set_xlim(0, max(vals) * 1.16)
    style_ax(ax)
    fig.savefig(TANKER_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_lpcd_chart():
    lp = sorted([t for t in EXTRA["lpcd"] if 50 < t[1] < 400 and "governorate" not in t[0].lower()],
                key=lambda t: t[1])
    names = [t[0] for t in lp]; vals = [t[1] for t in lp]
    cols = [ORANGE if n == "Adh Dhahirah" else "#9CC7E4" for n in names]
    fig, ax = plt.subplots(figsize=(7.0, 3.4))
    ax.barh(names, vals, color=cols, height=0.62, zorder=3)
    for i, v in enumerate(vals):
        bold = names[i] == "Adh Dhahirah"
        ax.text(v + 2, i, f"{v:.1f}", va="center", fontsize=8.6, color=INK, fontweight="bold" if bold else "normal")
    ax.set_xlabel("domestic consumption 2024 (l/c/d)", fontsize=9, color=INK)
    ax.tick_params(axis="y", labelsize=8.6)
    ax.set_xlim(0, max(vals) * 1.15)
    style_ax(ax)
    fig.savefig(LPCD_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

def make_wwg_chart():
    series = [("IBRI", BLUE), ("AD DARIZ", ORANGE), ("AL ARAQI", "#009E73"), ("TOTAL", "#444444")]
    fig, ax = plt.subplots(figsize=(7.2, 3.8))
    for name, col in series:
        d = EXTRA["wwg"][name]
        yrs = sorted(int(y) for y in d)
        vals = [d[str(y)] / 1000 for y in yrs]
        ls = "--" if name == "TOTAL" else "-"
        ax.plot(yrs, vals, color=col, lw=2.0, ls=ls, zorder=3)
        lbl = "Project TOTAL" if name == "TOTAL" else name
        dy = {"AD DARIZ": 9, "AL ARAQI": -9}.get(name, 0)
        ax.annotate(lbl, xy=(yrs[-1], vals[-1]), xytext=(4, dy), textcoords="offset points",
                    fontsize=8.6, color=col, fontweight="bold", va="center")
    ax.set_xlim(2023, 2118)
    ax.set_xticks(range(2030, 2101, 10))
    ax.set_xlabel("year", fontsize=9, color=INK)
    ax.set_ylabel("WW generation (1000 m3/d, incl. +20% weekly peak)", fontsize=8.5, color=INK)
    style_ax(ax)
    fig.savefig(WWG_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

PFCHART = os.path.join(IMG, "t01_pf.png")
def make_pf():
    q = np.logspace(0, 3, 200)              # L/s
    pf_pel = np.clip(1.5 + 1 / np.sqrt(q), None, 5.0)
    q_mld = q * 0.0864
    pf_mer = 2.65 * q_mld ** (-0.121)
    fig, ax = plt.subplots(figsize=(7.0, 3.6))
    ax.semilogx(q, pf_pel, color=BLUE, lw=2.2, zorder=3)
    ax.semilogx(q, pf_mer, color=ORANGE, lw=2.2, zorder=3)
    ax.axhline(5.0, color="#999999", lw=1.1, ls=":")
    ax.text(1.15, 5.08, "cap PF = 5.0 (G201-p72)", fontsize=8.5, color="#777777")
    ax.text(300, 1.62, "Peltier (peak hour)", fontsize=9.5, color=BLUE, fontweight="bold")
    ax.text(300, 2.35, "Merrimack (peak day)", fontsize=9.5, color=ORANGE, fontweight="bold")
    ax.plot([20], [1.72], "o", color=BLUE, ms=6)
    ax.plot([20], [2.48], "o", color=ORANGE, ms=6)
    ax.annotate("worked example\nQadf = 20 L/s", xy=(20, 1.72), xytext=(33, 1.1), fontsize=8.8, color=INK,
                arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.set_xlabel("average flow Qadf (L/s, log scale)", fontsize=9, color=INK)
    ax.set_ylabel("peak factor PF", fontsize=9, color=INK)
    ax.set_ylim(0.8, 5.6)
    style_ax(ax)
    fig.savefig(PFCHART, dpi=200, bbox_inches="tight"); plt.close(fig)

# ---------- Colebrook-White hydraulics (Section 14; ks = 1.5 mm G203-p24/28, nu = 1.141e-6 G203-p25) ----------
import math
CW_KS = 1.5e-3; CW_NU = 1.141e-6; CW_G = 9.81
def cw_v(Dh, S):
    """Full-bore Colebrook-White velocity (m/s) for hydraulic diameter Dh (m) at slope S (m/m)."""
    root = math.sqrt(2 * CW_G * Dh * S)
    return -2.0 * root * math.log10(CW_KS / (3.71 * Dh) + 2.51 * CW_NU / (Dh * root))
def cw_partial(D, S, dD):
    """(velocity, area, discharge) at relative depth d/D via circular-segment geometry."""
    th = 2 * math.acos(1 - 2 * dD)
    A = D * D / 8 * (th - math.sin(th))
    R = A / (th * D / 2)
    v = cw_v(4 * R, S)
    return v, A, v * A
TAB11 = [(200, 5.00), (250, 3.75), (315, 2.70), (400, 2.05),
         (500, 1.55), (600, 1.25), (700, 1.00), (800, 0.85), (900, 0.75)]

CW_PARTIAL_CHART = os.path.join(IMG, "t01_cw_partial.png")
def make_cw_partial():
    D, S = 0.315, 0.004
    vfull = cw_v(D, S); qfull = vfull * math.pi * D * D / 4
    dd = np.linspace(0.02, 0.999, 400)
    vr, qr = [], []
    for x in dd:
        v, _, q = cw_partial(D, S, float(x))
        vr.append(v / vfull); qr.append(q / qfull)
    fig, ax = plt.subplots(figsize=(7.0, 3.8))
    ax.plot(qr, dd, color=BLUE, lw=2.2, zorder=3)
    ax.plot(vr, dd, color=ORANGE, lw=2.2, zorder=3)
    ax.axhline(0.65, color="#777777", lw=1.1, ls="--")
    ax.axhline(0.50, color="#AAAAAA", lw=1.1, ls=":")
    ax.text(0.03, 0.665, "d/D limit 0.65 (D ≤ 350 mm, G203-p27 Tab 10)", fontsize=8.3, color="#555555")
    ax.text(0.03, 0.515, "d/D limit 0.50 (D > 350 mm)", fontsize=8.3, color="#888888")
    ax.text(0.80, 0.30, "Q / Qfull", fontsize=9.5, color=BLUE, fontweight="bold")
    ax.text(1.02, 0.16, "V / Vfull", fontsize=9.5, color=ORANGE, fontweight="bold")
    ax.plot([0.639], [0.581], "o", color=BLUE, ms=6)
    ax.annotate("worked example:\n45 L/s in DN315 @ 4.0 mm/m", xy=(0.639, 0.581), xytext=(0.28, 0.78),
                fontsize=8.6, color=INK, arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.annotate("V peaks at d/D ≈ 0.81", xy=(max(vr), 0.81), xytext=(0.62, 0.93), fontsize=8.6, color=INK,
                arrowprops=dict(arrowstyle="->", color="#777777"))
    ax.set_xlim(0, 1.25); ax.set_ylim(0, 1.0)
    ax.set_xlabel("ratio to full-bore value", fontsize=9, color=INK)
    ax.set_ylabel("relative depth d/D", fontsize=9, color=INK)
    style_ax(ax)
    fig.savefig(CW_PARTIAL_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

CW_TAB11_CHART = os.path.join(IMG, "t01_cw_tab11.png")
def make_cw_tab11():
    dns = [t[0] for t in TAB11]; smins = [t[1] for t in TAB11]
    vs = [cw_v(dn / 1000.0, s / 1000.0) for dn, s in TAB11]
    x = np.arange(len(dns))
    fig, ax = plt.subplots(figsize=(7.0, 3.2))
    ax.axhline(0.75, color=ORANGE, lw=1.6, ls="--", zorder=2)
    ax.text(0.05, 0.7535, "self-cleansing target 0.75 m/s (G203-p26)", fontsize=8.6, color=ORANGE)
    ax.plot(x, vs, "o", color=BLUE, ms=7, zorder=3)
    for i, v in enumerate(vs):
        ax.text(x[i], v + 0.004, f"{v:.3f}", ha="center", fontsize=8.0, color=BLUE)
        ax.text(x[i], 0.712, f"{smins[i]:.2f}", ha="center", fontsize=7.6, color="#777777")
    ax.text(-0.45, 0.7045, "Table 11 gradient (mm/m):", fontsize=7.6, color="#777777", ha="left")
    ax.set_xticks(x); ax.set_xticklabels([f"DN{d}" for d in dns], fontsize=8.6)
    ax.set_ylim(0.70, 0.79)
    ax.set_ylabel("CW full-bore velocity (m/s)", fontsize=9, color=INK)
    style_ax(ax)
    fig.savefig(CW_TAB11_CHART, dpi=200, bbox_inches="tight"); plt.close(fig)

make_chain(); make_diurnal(); make_pop(); make_pf(); make_cw_partial(); make_cw_tab11()
if POP_SERIES: make_pop_settlements()
make_demand_buildup(); make_return_chart()
if EXTRA: make_ramp_chart(); make_tanker_chart(); make_lpcd_chart(); make_wwg_chart()
route_chart(ROUTE_CENSUS, "Route A — Census projection (dated horizons 2030 / 2055)", [
    ("NCSI wilayat population series", "official census + forecast, Ibri wilayat, from 2020 onward (G201-p58-59)"),
    ("Project the growth", "apply the NCSI forecast (or fitted growth rate) at 5-year intervals to each model year"),
    ("Disaggregate to settlements", "split the wilayat total using the settlement shares of the latest census (G201-p58)"),
    ("Sub-settlement split where needed", "pro-rata by electricity account counts within the area, provided by NWS (G201-p58)"),
    ("Population per zone and model year", "input to water demand (Step 2) for 2030 / 2055 horizons"),
])
route_chart(ROUTE_PLOTS, "Route B — Plot count x occupancy (ultimate / saturated horizon)", [
    ("Cadastral plots (MoHUP)", "count plots by land-use typology inside each zone (G201-p58)"),
    ("x average properties per plot", "how many dwellings a developed plot carries at build-out (subdivision checked)"),
    ("x occupancy rate = 5", "people per property; set 2026-08-19. Properties per plot are counted from electricity accounts, not assumed (G201-p58-59)"),
    ("Saturation population per zone", "the ceiling when every plot is developed - independent of growth rate"),
], color="#8A5A00", fill="#FBF3E2")

# ================================================================= DOCUMENT
doc = Document(SAMPLE)
body = doc.element.body
sect = body.find(qn("w:sectPr"))
for el in list(body):
    if el is not sect:
        body.remove(el)

def patch_part(part):
    xml = part._element
    first = True
    for t in xml.iter(qn("w:t")):
        if first:
            t.text = PROJ + " — " + SUB
            first = False
        else:
            t.text = ""
for s in doc.sections:
    for hd in (s.header, s.first_page_header, s.even_page_header):
        if hd is not None: patch_part(hd)

def para(text, style=None, align=None, bold=None, size=None, color=None, space_after=6, italic=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if bold is not None: r.bold = bold
    if italic is not None: r.italic = italic
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p
def H(level, text):
    return doc.add_paragraph(text, style=f"Heading {level}")
_styles = {s_.name for s_ in doc.styles}
def bullet(text, bold_lead=None):
    if "List Bullet" in _styles:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run("• ")
    if bold_lead:
        r = p.add_run(bold_lead); r.bold = True
    p.add_run(text)
    return p

FIGS, TABS = [], []
def fig_caption(text):
    n = len(FIGS) + 1
    label = f"Figure {n} — {text}"
    FIGS.append(label)
    try: p = doc.add_paragraph(label, style="Caption")
    except KeyError: p = para(label, bold=True, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p
def tab_caption(text):
    n = len(TABS) + 1
    label = f"Table {n} — {text}"
    TABS.append(label)
    try: p = doc.add_paragraph(label, style="Caption")
    except KeyError: p = para(label, bold=True, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p
def pic(path, w=6.0):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)
def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    try: t.style = "Table Grid"
    except KeyError: pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(font)
        shade(c, "D9E2F3")
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(font)
    if widths:
        for i, w in enumerate(widths):
            for r_ in t.rows:
                r_.cells[i].width = Inches(w)
    return t
def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def toc():
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "Right-click and Update Field to refresh."
    r.append(t); fld.append(r); p._p.append(fld)

# ---------- OMML native equations ----------
NSDECL = ('xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
          'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')
def mr(t, sty=None):
    rpr = '<m:rPr><m:sty m:val="%s"/></m:rPr>' % sty if sty else ""
    return f'<m:r>{rpr}<w:rPr><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/></w:rPr><m:t xml:space="preserve">{t}</m:t></m:r>'
def frac(num, den):
    return f'<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>'
def sqrt(e):
    return f'<m:rad><m:radPr><m:degHide m:val="1"/></m:radPr><m:deg/><m:e>{e}</m:e></m:rad>'
def ssub(base, sub):
    return f'<m:sSub><m:e>{base}</m:e><m:sub>{sub}</m:sub></m:sSub>'
def ssup(base, sup):
    return f'<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>'
def equation(inner):
    xml = f'<w:p {NSDECL}><w:pPr><w:jc w:val="center"/><w:spacing w:after="120"/></w:pPr><m:oMathPara><m:oMath>{inner}</m:oMath></m:oMathPara></w:p>'
    sect.addprevious(parse_xml(xml))

Q = lambda s: ssub(mr("Q"), mr(s))
q_ = lambda s: ssub(mr("q"), mr(s))

# ================= COVER =================
para("", space_after=60)
para("Sultanate of Oman", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
para("Nama Water Services Company SAOC (NWS / OWWSC)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
para("", space_after=40)
try:
    para(PROJ, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
except KeyError:
    para(PROJ, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
para("Tender No. T/2719110/2025", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("", space_after=40)
para(SUB, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=20, color="0000CC")
para("Revision 4 — August 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para("Rev 4 (2026-08-19): properties per plot are now counted from electricity accounts and the occupancy rate is fixed at 5 people per property, so Route B rests on measured numbers. Rev 3 added the pipe hydraulics chapter (Section 14). Rev 2 addressed review comments on Rev 0 and Rev 1.", align=WD_ALIGN_PARAGRAPH.CENTER, size=9, italic=True)
para("", space_after=80)
para("Renardet S.A. & Partners Consulting Engineers", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
para("Project No. 2621", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
pagebreak()

# ================= ABBREVIATIONS =================
para("Abbreviations and Symbols", bold=True, size=16, color="0000CC")
table(["Abbreviation", "Meaning"], [
 ["AAF", "Annual Average Flow — the sewage flow averaged over a full year (= Qadf in this tutorial)"],
 ["ADF / Qadf", "Average Daily Flow — daily flow averaged over the year; the basic planning flow"],
 ["MDF", "Maximum Day Flow — the flow of the highest-flow day of the year"],
 ["PHF", "Peak Hour Flow — the flow of the highest hour; sizes all hydraulic pass-through elements"],
 ["PF", "Peak Factor — ratio converting an average flow into a peak flow"],
 ["BOD5", "Biochemical Oxygen Demand (5-day) — oxygen consumed by bacteria degrading the organic matter"],
 ["COD", "Chemical Oxygen Demand — oxygen equivalent of ALL oxidisable matter (chemical test)"],
 ["TSS", "Total Suspended Solids — particulate matter carried by the sewage"],
 ["LPCD", "Litres Per Capita per Day"],
 ["NCSI", "National Centre for Statistics and Information (Oman census authority)"],
 ["IMP", "Integrated Master Plan — NWS's national water/wastewater master plan (IMP 2024 edition)"],
 ["NWS", "Nama Water Services (the Client; also OWWSC)"],
 ["MoHUP", "Ministry of Housing and Urban Planning (cadastral plot data)"],
 ["STP", "Sewage Treatment Plant"],
 ["TSE / TE", "Treated Sewage Effluent / Treated Effluent — the reclaimed water product of the STP"],
 ["SLS", "Sewage Lifting Station"],
 ["OR", "Occupancy Rate — persons per housing unit"],
 ["G201 / G202 / G203", "NWS guidelines PAM-GUD-201 (General), -202 (Water & TSE), -203 (Wastewater)"],
], widths=[1.2, 5.3], font=9)
pagebreak()

# ================= EXECUTIVE SUMMARY =================
H(1, "Executive Summary")
para("This tutorial teaches, from first principles, how the sewage flows and pollution loads of the Ibri project are calculated: what each parameter means, where its value comes from, why the calculation is structured the way it is, and how the results are used to size the collection network and the sewage treatment plant (STP). It is written for a reader new to wastewater flow estimation; every value is cited to the NWS guidelines (G203, G201, G202, by page) or to the project Inception Report R0 demand model, and every equation is explained before it is used.")
para("The calculation is a nine-step chain. It is summarised here in full; each step has a dedicated section in the body of the report.")
bullet("Population is estimated by two complementary routes. ", bold_lead="Step 1 — Population (Section 5). ")
para("Route A projects the official NCSI census series for Ibri wilayat (183,564 inhabitants in 2024, growth declining from about 3.0 to 1.3 percent per year in the NCSI forecast to 2033) and disaggregates it to settlements using census shares — and below settlement level, pro-rata by electricity account counts, as the guideline directs. Route A produces the populations of the dated model years (2030, 2055). Route B multiplies the count of cadastral plots by the average number of properties per plot and by the occupancy rate. Since August 2026 both of those are measured rather than assumed: each electricity account is one property, so the properties on a plot are counted directly from the account layer (33,970 accounts across the wilayat, 1.4 properties per plot on average), and the occupancy rate is set at 5 people per property. Counting domestic accounts including the 'additional account' type gives 7.28 people per plot, against 6.96 obtained independently by dividing the NCSI population by the built plots — the two agree, which is why additional accounts are treated as separate dwellings. Route B produces the saturation (ultimate) population — the ceiling reached when every plot is developed, which is independent of how fast growth occurs. The network pipes, which live 50 years and cannot be economically re-laid, are sized for the Route B ceiling; the STP, which is built in phases, is sized for the Route A dated horizons. This is how the two routes are reconciled.")
bullet("Total water demand is built from the measured domestic consumption of Adh Dhahirah Governorate — 164 litres per capita per day (l/c/d) — plus governorate-average allowances of 22 percent for non-domestic (commercial) and 14 percent for governmental consumption, giving 223 l/c/d. These percentage allowances are a planning-level bookkeeping device: they distribute the governorate's measured non-domestic water use over the resident population so that totals are correct even before land-use-specific data exist. In the hydraulic model, large individual facilities are instead loaded at their own nodes using per-facility unit rates (litres per pupil, per bed, per m2).", bold_lead="Step 2 — Water demand (Section 6). ")
bullet("Only part of the supplied water reaches the sewer: 85 percent of domestic (and tankered) consumption and 54 percent of non-domestic and governmental consumption; the remainder is lost to irrigation, evaporative cooling and other consumptive uses. The result for Ibri is a wastewater generation of about 171 l/c/d.", bold_lead="Step 3 — Wastewater return (Section 7). ")
bullet("Two flows join the sewage: infiltration of groundwater and soil moisture through pipe joints and manholes (720 litres per day per kilometre of new sewer per the guideline; the R0 model provisionally uses the more conservative 10 percent of wastewater flow), and tankered sewage from properties not yet connected to the network (about 17 percent of STP inflow nationally today, to be phased out as coverage reaches 100 percent). A recent site visit found tankers arriving from camps up to 150 km away — beyond any assumption in the scope — which is flagged as a project risk.", bold_lead="Step 4 — Additions (Section 8). ")
bullet("The average daily flow Qadf is the population times the per-capita wastewater generation, plus infiltration and tanker deliveries. Qadf is the backbone quantity of the whole design: the STP biology, the annual loads and the phasing analysis all rest on it.", bold_lead="Step 5 — Average flow (Section 9). ")
bullet("Sewage flow varies strongly over the day, so pipes must carry the peak, not the average. The peak factor (PF) converts average to peak flow; it decreases as catchments grow because many households' peaks do not coincide. Two empirical formulas are permitted — Peltier (a peak-hour factor, the current IMP 2024 method) and Merrimack (a peak-day regression). They answer slightly different questions and were fitted to different data, so they disagree (1.72 vs 2.48 at the worked-example flow); the binding choice is to be confirmed with NWS at kickoff.", bold_lead="Step 6 — Peak flows (Section 10). ")
bullet("Pollution is quantified independently of water use: each person contributes a fixed daily mass of organic matter (at least 60 g of BOD5 and 80 g of TSS per day) regardless of how much water dilutes it. Loads (kg/d) size the biological treatment; concentrations (mg/l) are always derived as load divided by flow — Oman's low water use makes the sewage concentrated, around 300-400 mg/l BOD5.", bold_lead="Step 7 — Organic loads (Section 11). ")
bullet("Each design element is sized by a specific flow: gravity pipes and pumping stations by the peak-hour flow, the STP hydraulic pass-through by PHF, the biological process by AAF plus loads. The STP incoming flow adds a 10 percent operational allowance; the plant category (above/below 20,000 m3/d) determines the STP scope route per the Terms of Reference.", bold_lead="Steps 8-9 — Design flows and STP (Sections 12-13). ")
bullet("The STP returns 95 percent of its inflow as treated sewage effluent (TSE), which feeds the TE network design, and produces sludge (about 0.25 kg per m3 treated) which is thickened, dewatered and directed to reuse or disposal under the sludge management strategy required by the scope.", bold_lead="Outputs (Section 13). ")
bullet("New in Rev 3: how a pipe actually carries the peak flow. The Colebrook-White equation with the guideline's fixed inputs (ks = 1.5 mm, nu = 1.141e-6 m2/s), partial-full flow and the d/D limits, the demonstration that the Table 11 minimum gradients are simply Colebrook-White solved at the self-cleansing velocity 0.75 m/s, a worked pipe-sizing example, and the tractive-force minimum gradient for network heads. This is the method the W5 network-design pipeline implements.", bold_lead="Pipe hydraulics (Section 14). ")
para("A fully worked numerical example (Section 15) carries a hypothetical settlement of 10,000 persons through all nine steps, and a reconciliation register (Section 16) lists the four points where the R0 model and the guidelines differ (infiltration basis, weekly peak, tanker catchment radius, sludge rate) — all flagged for confirmation with NWS at the kickoff stage. Conclusions and recommendations close the report.")
pagebreak()

# ================= TOC / LOF / LOT =================
para("Table of Contents", bold=True, size=16, color="0000CC")
toc()
para("", space_after=10)
para("List of Figures", bold=True, size=14, color="0000CC")
LOF_MARK = doc.add_paragraph()
para("List of Tables", bold=True, size=14, color="0000CC")
LOT_MARK = doc.add_paragraph()
pagebreak()

# ================= 1 INTRODUCTION =================
H(1, "Introduction")
H(2, "Purpose and how to read this tutorial")
para("This document is a teaching text, not a design report: its goal is that the reader can afterwards perform and check every calculation independently. Each section first explains the concept in plain terms — what the quantity is, why it exists, what would go wrong without it — then gives the governing values and equations with their sources, and finally shows the numbers at work in the worked example of Section 15. No parameter is used without being defined, and no equation appears without its purpose stated.")
H(2, "Reference documents and citation convention")
table(["Citation", "Document"], [
 ["G203-p##", "PAM-GUD-203 — Wastewater Design Guidelines v1.0 (Rev 01), page ##"],
 ["G201-p##", "PAM-GUD-201 — General Design Guidelines v1.0 (Rev 01), page ##"],
 ["G202-p##", "PAM-GUD-202 — Potable Water & TSE Design Guidelines v1.0 (Rev 01), page ##"],
 ["TOR", "Terms of Reference, Section 03, Tender T/2719110/2025"],
 ["R0", "Inception Report R0 demand workbook (Ibri Sewer Demand R0, August 2026)"],
], widths=[1.1, 5.4])
tab_caption("Reference documents and citation convention")
para("Where the guidelines and the R0 model differ, both values are given and the difference is flagged for confirmation with NWS (Section 16).")

# ================= 2 KEY CONCEPTS =================
H(1, "Key Concepts: How Sewage Flow Behaves")
para("Before any calculation, three facts about sewage flow must be understood, because the whole structure of the method follows from them.")
bullet("Sewage flow follows human activity over the day. It is lowest in the small hours of the night (when what remains in the pipe is mostly infiltration), rises steeply with the morning routine, and shows a second evening peak. Figure 1 shows a typical (illustrative) diurnal pattern.", bold_lead="It varies over the day. ")
bullet("Weekends, holidays and seasons shift water use; over a year, the highest-flow day can be well above the average day.", bold_lead="It varies over the year. ")
bullet("A single house may momentarily discharge many times its average; a city of 100,000 never does, because individual peaks do not coincide. This damping with size is why peak factors decrease as flow accumulates downstream, and why small upstream sewers need proportionally more capacity headroom than large trunk mains.", bold_lead="The bigger the catchment, the smoother the flow. ")
pic(DIURNAL, w=5.9)
fig_caption("Typical diurnal sewage flow pattern (illustrative), with the average and peak-hour levels marked")
para("From this behaviour the guidelines define three reference flows (G203-p65-66), each with a distinct design role:")
table(["Flow", "Definition", "What it is used for"], [
 ["AAF (= Qadf)", "Annual Average Flow: total annual volume / 365 days. In this tutorial written Qadf (average daily flow).", "The planning backbone: STP biological sizing, annual loads, phasing, energy and cost estimates"],
 ["MDF", "Maximum Day Flow: the highest single-day volume of the year.", "Storage, emergency lagoons, day-scale process checks"],
 ["PHF", "Peak Hour Flow: the highest hourly flow of the year.", "Everything sewage physically passes through: pipes, pumps, screens, channels, STP headworks"],
], widths=[1.1, 2.7, 2.7], font=9)
tab_caption("The three reference flows of G203-p65-66 and their design roles")
para("The relationship is always AAF < MDF < PHF. The calculation chain of this tutorial first establishes Qadf (Steps 1-5), then converts it to the peaks (Step 6).")

# ================= 3 CHAIN =================
H(1, "The Calculation Chain at a Glance")
para("The nine steps of Figure 2 form a strict chain: each step consumes the previous step's output. Steps 1-5 build the average flow; Step 6 converts it to peak flows; Step 7 establishes the pollution loads on a parallel track (they depend on population, not on water); Steps 8-9 combine flows and loads at the treatment plant.")
pic(CHAIN, w=5.6)
fig_caption("The sewage flow and load calculation chain (Steps 1-9)")
pagebreak()

# ================= 4 STEP 1 POPULATION =================
H(1, "Step 1 — Population")
H(2, "Why population comes first")
para("Every quantity in this tutorial — water, sewage, BOD, sludge — is generated by people. Population is therefore the root input, and any error in it propagates proportionally through the entire chain: a 10 percent error in population is a 10 percent error in STP capacity. The guidelines (G201-p57) recognise three approaches: figures provided directly by a developer (for defined developments), calculation from NCSI forecast data, and calculation from plots, property types and occupancy rates. For a whole-town project like Ibri the last two are used together, as Routes A and B below; developer figures apply only to specific master-planned schemes if NWS provides them.")
H(2, "Route A — census projection (for the dated model years)")
para("Route A answers: how many people will actually live here in 2030, in 2055? It starts from the official NCSI population series and works downward in scale:")
bullet("The NCSI series for Ibri wilayat is the authoritative starting point (G201-p58-59). The current figures, taken from the R0 workbook, are shown in Table 4 and Figure 3.", bold_lead="Start from the wilayat series. ")
bullet("NCSI publishes forecasts; where projection beyond them is needed, a growth rate fitted to the series is applied at the 5-year intervals required by the TOR (start year, 2030, 2055). For this project the route is already implemented: the R0 workbook tab 'Project Pop Settlements' carries the projected population of all 50 project settlements, year by year, out to 2100 (Appendix B).", bold_lead="Project the growth. ")
bullet("The wilayat total is split over settlements in proportion to their shares in the latest census (G201-p58) — a settlement holding 4 percent of the census population receives 4 percent of every projected total.", bold_lead="Disaggregate to settlements. ")
bullet("Where a project zone covers only part of a settlement, the guideline directs a pro-rata split by the number of electricity accounts in the area, obtained from NWS (G201-p58). Electricity accounts are a reliable proxy for occupied dwellings because effectively every occupied dwelling has a metered connection — this is also why the TOR (p12) lists electricity data among the required flow-assessment inputs.", bold_lead="Split below settlement level by electricity accounts. ")
pic(ROUTE_CENSUS, w=5.9)
fig_caption("Route A — from NCSI wilayat series to zone populations for the dated horizons")
table(["Year", "2021", "2022", "2023", "2024", "2026", "2028", "2030", "2032", "2033"], [
 ["Population", "168,409", "173,418", "178,477", "183,564", "193,116", "203,659", "213,637", "224,840", "227,736"],
 ["Growth (%/yr)", "—", "2.97", "2.92", "2.85", "2.74", "2.68", "2.19", "2.58", "1.29"],
], widths=[1.05, 0.62, 0.62, 0.62, 0.62, 0.62, 0.62, 0.62, 0.62, 0.62], font=8)
tab_caption("NCSI population series for Ibri wilayat (from the R0 workbook; selected years)")
pic(POPCHART, w=5.9)
fig_caption("NCSI population series and forecast for Ibri wilayat, 2021-2033")
if POP_SERIES:
    para(f"At project level, the R0 workbook projects each of the 50 settlements to 2100. Figure 5 shows the three largest (IBRI town, AD DARIZ, AL ARAQI) and the project total: from {(_sv('TOTAL', 2024)/1000):.0f} thousand (2024) the project population reaches {(_sv('TOTAL', 2055)/1000):.0f} thousand by 2055 and {(_sv('TOTAL', 2100)/1000):.0f} thousand by 2100 — the R0 projection grows throughout, i.e. it contains no saturation ceiling of its own (see the discussion of Section 5.4).")
    pic(SETTLCHART, w=6.0)
    fig_caption("R0 projected populations to 2100 — IBRI, AD DARIZ, AL ARAQI and project total (workbook tab 'Project Pop Settlements')")
H(2, "Route B — plots x occupancy (for the ultimate horizon)")
para("Route B answers a different question: how many people can this area hold when it is fully built? It works upward from the cadastre:")
equation(mr("Population") + mr(" = ") + ssub(mr("N"), mr("plots")) + mr(" × ") + ssub(mr("n"), mr("prop")) + mr(" × ") + mr("OR"))
para("where N(plots) is the number of plots of a given typology inside the zone, n(prop) is the average number of properties (dwellings) each such plot carries at build-out, and OR is the occupancy rate (G201-p58):")
equation(mr("OR") + mr(" = ") + frac(mr("Population"), mr("Housing units")))
para("Both OR inputs come from the NCSI data portal, which publishes population and housing-unit counts from 2020 at governorate and wilayat level (G201-p59). The guideline adds two cautions:")
bullet("check plot subdivision, so that the effective number of housing units created is not understated when large plots split into several dwellings (G201-p59).", bold_lead="Subdivision: ")
bullet("temper the assumption that all plots fill up — build-out takes decades, so the saturation population is a ceiling, not a date (G201-p59).", bold_lead="Development speed: ")
pic(ROUTE_PLOTS, w=5.9)
fig_caption("Route B — from cadastral plots to the saturation population")
H(2, "How the two routes fit together — saturation vs dated horizons")
para("The two routes serve different design decisions, and understanding their relationship is essential:")
bullet("The master-plan plots define the maximum development the planning framework allows. Multiplying all plots by properties-per-plot and occupancy gives the saturation population — a ceiling, not a forecast. It says nothing about when that ceiling is reached.")
bullet("The census projection gives the population trajectory in time — but knows nothing about the physical ceiling; extrapolated far enough it would eventually exceed what the plots can hold.")
bullet("The two are reconciled by using each where it is authoritative: buried pipes have a 50-year design life (G201-p57) and cannot economically be re-laid, so the network is dimensioned against the long-term (saturation) flows. The STP is modular and expanded in phases, so each phase is sized for a Route A dated horizon (2030, 2055).")
bullet("The TOR's design horizon rule — completion + 25 years or ultimate (saturated) — has two possible readings, and which applies depends on which event comes first. If saturation is reached before completion + 25 years, saturation is the design condition. If — as the client appears to expect, and the R0 projections support — saturation lies beyond completion + 25 years (i.e. after roughly 2055), then the contractual design flow is that of the governing 5-year interval projection year (populations are computed for every year, but the design horizon snaps to the TOR's 5-year grid), and the saturation flow serves as the long-term capacity check for the trunk sewers. The binding reading shall be confirmed with NWS at kickoff.")
bullet("Consistency check: zone-by-zone, the Route A trajectory must be compared against the Route B ceiling. The current numbers show why this matters: the R0 projection for the project area grows continuously to about 691 thousand by 2100 with no built-in ceiling, while the provisional plot-based saturation (at the assumed occupancy of 6.0) is far lower — the two cross somewhere in mid-century. Either the occupancy/properties-per-plot assumptions understate the ceiling, or the projection overshoots the developable land. Resolving this — with NCSI housing data and MoHUP planned plots — is a first-priority kickoff task, because it decides which quantity governs the network.")
H(2, "Future plots and stub-outs — how undeveloped land enters the design")
para("Within the project boundary many plots are only planned polygons — subdivided but empty land. They generate no sewage today, yet the TOR requires that all plots, built, open and under construction, be provided with connections. The design handles them as follows:")
bullet("every plot, including future ones, contributes its saturation population to the flows used to dimension the pipes. This is precisely why Route B counts plots rather than people: an empty plot is a future flow with a known magnitude and location.", bold_lead="Pipes are dimensioned for the long-term flow including future plots: ")
bullet("at the frontage of each future plot or undeveloped sub-area, a capped connection (stub-out) or capped lateral is built with the main sewer, sized for that area's future flow. When the plot develops, it connects without excavating the road again. Downstream of every stub-out, the pipes are dimensioned to carry the future inflow — so yes, the reader's understanding is correct: a stub-out is a commitment that flow will one day enter there, and the downstream system is designed to accept it.", bold_lead="Stub-outs reserve the connection: ")
bullet("population and flow projections (Route A) distribute the dated-year population over the settlements as they exist and grow; future plots fill progressively toward saturation. The dated-year flows therefore drive the phased facilities (STP modules, pumping equipment), not the pipe diameters.", bold_lead="Projections fill the plots over time: ")
bullet("a sewer dimensioned for future flows runs nearly empty in its early years, and the risk is sediment deposition, not capacity. This is why the minimum self-cleansing criteria (0.75 m/s at peak, or the tractive-force check at network heads, G203-p26-27) must be verified at the start-year flows, not only at the horizon — the standard early-years check of every staged network.", bold_lead="The early-years check is the flip side: ")
H(2, "Is the available plot data ready for this?")
para("Partly. The MoHUP cadastral layer for Ibri contains 61,272 plot polygons with land-use classification and a built-type attribute, which is sufficient for counting and classifying plots per zone (Route B's first input). Two of the three factors are, however, not yet supported by data: the average properties per plot (requires the subdivision check of G201-p59) and the occupancy rate (requires NCSI housing-unit counts, which are not included in the R0 package — its workbook carries population only). Both are on the kickoff data request; until they arrive, screening uses a tagged assumption (OR = 6.0) and results are reported as scaling linearly with it.")

# ================= 5 STEP 2 WATER =================
H(1, "Step 2 — Water Demand")
H(2, "The concept: sewage begins as drinking water")
para("Nearly all sewage is water that was first supplied through the potable network (or by water tanker), used, and discharged to the drain. Wastewater estimation therefore starts by establishing how much water each person's presence causes to be consumed — not only at home (domestic), but also in the shops, offices, schools and government buildings that exist because the population exists.")
H(2, "The per-capita method and why the components are added")
para("The domestic consumption of Adh Dhahirah Governorate is 164 l/c/d — a measured value from the IMP 2024 baseline (G201-p60), which the R0 model independently reproduces (163.5 l/c/d) from 2021-24 billing records. Non-domestic and governmental consumption are then expressed as governorate-average ratios of the domestic figure: +22 percent and +14 percent respectively (G201-p60-61).")
para("A reader may reasonably object: commercial and governmental demand belongs to specific land-use plots, not to residents — why add it per capita? The answer is that the ratios are an aggregate bookkeeping device, and it is worth being precise about what they do and do not claim:")
bullet("At governorate scale, the total non-domestic and governmental water use is measured (it is in the billing data). Dividing it by the resident population and expressing it as a percentage of domestic use distributes that measured total over the people it ultimately serves. The town-wide totals are then correct by construction — even before any land-use analysis exists. This is the appropriate method at the planning stage, when the question is 'how much sewage will Ibri produce'.")
bullet("At model scale, the same consumption is instead assigned spatially: each commercial or institutional plot is loaded at its own network node, using the unit rates of G201-p61 Table 12 (130 l/pupil/d for schools, 650 l/bed/d for hospitals, 12.2 l/m2/d for shopping floor area, 93 l/employee/d for offices, etc.), with the occupancy of each facility taken from the GIS land-use data. This is the appropriate method when the question is 'which pipe must carry the hospital's discharge'.")
bullet("The two views must reconcile: the sum of the spatially-assigned non-domestic loads over the whole town should approximate the percentage allowance times the population. A large discrepancy signals either unrealistic facility occupancies or a non-domestic sector unlike the governorate average — and is investigated, not averaged away.")
table(["Component", "Basis", "Value (l/c/d)", "Source"], [
 ["Domestic", "measured, Adh Dhahirah", "164", "G201-p60; R0: 163.5 from billing"],
 ["Non-domestic", "+22% of domestic (governorate ratio)", "36.1", "G201-p60-61"],
 ["Governmental", "+14% of domestic (governorate ratio)", "23.0", "G201-p60"],
 ["Total water demand", "sum (aggregate planning view)", "223", "—"],
], widths=[1.7, 2.4, 1.2, 1.9], font=9)
tab_caption("Water demand build-up for Ibri (aggregate per-capita view)")
pic(DEMAND_BUILDUP, w=6.0)
fig_caption("Water demand build-up: how 223 l/c/d is assembled from its three components")
para("Appendix B.4 charts the measured domestic consumption of all eleven governorates — Adh Dhahirah's 163.5 l/c/d sits close to the national average, supporting the 164 l/c/d design value.")

# ================= 6 STEP 3 RETURN =================
H(1, "Step 3 — Wastewater Return")
H(2, "The concept: not all water comes back")
para("Part of the supplied water never reaches the sewer: garden and farm irrigation percolates or evaporates, evaporative (desert) coolers and cooling towers evaporate their feed, washing water is thrown on yards, mosques' ablution water often irrigates landscaping. The return rate is the measured fraction that does come back, and it differs by consumer type (G201-p71, Table 19):")
table(["Stream", "Return rate", "Why this value"], [
 ["Domestic (and tankered)", "85%", "most indoor use returns; ~15% lost to gardens, coolers, outdoor cleaning"],
 ["Non-domestic + governmental", "54%", "large landscaped compounds, cooling and process losses make commercial/institutional use far more consumptive"],
], widths=[2.2, 1.2, 3.6], font=9)
tab_caption("Return rates to the sewer by consumer stream (G201-p71 Table 19)")
para("The per-capita wastewater generation follows by applying each return rate to its own stream and summing the results. As in Step 2, the summation is the aggregate planning view — each stream keeps its own rate, and only the resulting sewage quantities (which do all end up in the same pipe) are added:")
equation(q_("ww") + mr(" = ") + q_("dom") + mr(" × 0.85 + (") + q_("nd") + mr(" + ") + q_("gov") + mr(") × 0.54"))
para("For Ibri: 164 × 0.85 + (36.1 + 23.0) × 0.54 = 139.4 + 31.9 ≈ 171 l/c/d of sewage per person. This single number now carries the whole demand side of the calculation. The figure below shows the two streams side by side — supplied water against what actually returns as sewage:")
pic(RETURN_CHART, w=5.9)
fig_caption("From water to sewage: supplied vs returned volumes per stream (l/c/d)")

# ================= 7 STEP 4 ADDITIONS =================
H(1, "Step 4 — Additions: Infiltration and Tankered Sewage")
H(2, "Infiltration — what it is and how it is estimated")
para("Sewers are not watertight. Groundwater and soil moisture enter through pipe joints, manhole walls, and — in older systems — cracks and illegal connections. This infiltration flows continuously, day and night (it is what remains in the pipe at 4 a.m., Figure 1), consumes pipe and treatment capacity, and dilutes the sewage. It cannot be zero, so it is budgeted explicitly (G201-p72):")
bullet("720 litres per day per kilometre of sewer — a linear allowance reflecting tight modern joints and (in Ibri's inland setting) deep groundwater.", bold_lead="New networks: ")
bullet("10% of the wastewater flow; up to 40% for coastal or high-groundwater systems.", bold_lead="Existing inland networks: ")
bullet("None. The system is strictly separate; rainfall does not enter the design flows (G201-p72).", bold_lead="Stormwater allowance: ")
para("The R0 model provisionally applies the 10 percent rule to selected settlements. For the new Ibri network this is roughly ten times the guideline's linear allowance — a conservative choice that is safe for STP capacity but not free: carried through to ultimate flows it adds several thousand m3/d of phantom flow to the STP sizing and dilutes the design concentrations. The recommended treatment (Section 16) is to carry the guideline value as the design basis, keep 10 percent as an upper sensitivity bound, and have NWS confirm the basis at kickoff.")
H(2, "Tankered sewage")
para("Properties not yet connected to a network discharge to holding tanks emptied by vacuum tankers ('yellow tankers'), which discharge at the STP's tanker reception facility. Tanker sewage is therefore part of the STP inflow from day one, while network flows ramp up as connections are made. Nationally, tankers deliver about 17 percent of STP inflow today; design coverage reaches 100 percent by the end of the planning period, phasing tanker deliveries out (G201-p73). Two project-specific points:")
bullet("The R0 model collects tankered flows from settlements within 25 km of the STP. This radius is an R0 assumption with no guideline source — it must not be mixed into design calculations until NWS confirms it (Section 16).")
bullet("A recent site visit found tankers delivering sewage from camps up to 150 km away. This is outside any assumption in the scope and materially affects the tanker reception sizing, the early-years flow balance and possibly the STP odour/septicity design (long-haul sewage arrives septic). It is flagged as a project risk requiring an NWS policy decision: which catchment is the Ibri STP obliged to accept?")

# ================= 8 STEP 5 AVERAGE =================
H(1, "Step 5 — Average Daily Flow")
para("The average daily flow assembles everything so far: people times per-capita sewage, plus the two additions of Step 4:")
equation(Q("adf") + mr(" = ") + frac(mr("P × ") + q_("ww"), mr("1000")) + mr(" + ") + Q("inf") + mr(" + ") + Q("tank"))
para("where P is the population served (capita), q(ww) the per-capita wastewater generation (l/c/d, = 171 for Ibri), Q(inf) the infiltration allowance and Q(tank) the tankered deliveries, all in m3/d. The division by 1000 converts litres to cubic metres.")
para("Qadf equals the AAF of Section 3 and is the backbone quantity of the design: every later flow is derived from it, the biological STP sizing rests on it, and the phasing analysis compares it across model years. In the worked example (Section 15) it evaluates to 1,731 m3/d, i.e. 20.0 L/s.")

# ================= 9 STEP 6 PEAKS =================
H(1, "Step 6 — Peak Flows")
H(2, "What a peak factor is and why it shrinks with catchment size")
para("Pipes and pumps must carry the worst hour, not the average day — a sewer sized for Qadf would surcharge every morning (Figure 2). The peak factor is the multiplier that converts the average flow into that governing peak. Because individual households' peaks do not coincide, the aggregate flow smooths out as the catchment grows: PF is large for a street (approaching the cap of 5), small for a city trunk (approaching about 1.6). All peak-factor formulas are therefore decreasing functions of flow. Peaking applies to the sewage component only — infiltration is a steady base flow and is added after peaking.")
H(2, "The two permitted formulas — and why they differ")
para("The guidelines permit two empirical formulas (G201-p71-72). They are not two estimates of the same number; they answer different questions:")
bullet("gives the ratio of the peak-hour flow to the average flow. Developed for the Oman IMP 2024, it is the current national method. Qm is the mean flow in L/s:", bold_lead="Peltier (peak-hour factor) ")
equation(mr("PF") + mr(" = 1.5 + ") + frac(mr("1"), sqrt(Q("m"))))
bullet("is a regression fitted to observed maximum-day flows (US practice, valid above ~100 connected properties); it returns the peak-day flow directly, both flows in Ml/d:", bold_lead="Merrimack (peak-day flow) ")
equation(Q("pdf") + mr(" = 2.65 × ") + ssup(mr("(") + Q("adf") + mr(")"), mr("0.879")))
para("Both are capped by the absolute limit PF ≤ 5.0 on the hourly factor (G201-p72). Figure 9 plots the two over the practically relevant flow range.")
pic(PFCHART, w=5.9)
fig_caption("Peltier and Merrimack peak factors versus average flow, with the PF = 5.0 cap and the worked example marked")
para("Why do they disagree (1.72 vs 2.48 at 20 L/s)? Three reasons, all visible in Figure 9:")
bullet("Peltier's 1.5 asymptote is a peak-hour ratio for large flows; Merrimack tracks the maximum day, and its implied factor (2.65 × Qadf to the power −0.121) stays higher across the range. A peak day and a peak hour are different events.", bold_lead="Different peak definitions. ")
bullet("Peltier was fitted to Omani IMP data; Merrimack to New England systems with different habits, appliances and infiltration behaviour. Empirical regressions carry their calibration data with them.", bold_lead="Different source data. ")
bullet("Merrimack's exponent makes its factor fall very slowly with size, so the divergence grows for large flows.", bold_lead="Different damping. ")
H(2, "Which formula for which element — and where Qm comes from")
para("A common shorthand is 'Peltier for the pipes, Merrimack for the STP'. That is close, but one refinement is needed:")
bullet("Network hydraulic sizing (pipes, pump stations) — Peltier. These elements are governed by the peak hour, which is what Peltier estimates; it is also the current IMP 2024 method.")
bullet("STP — both, in different places. The plant's hydraulic pass-through (headworks, screens, channels) is also sized on the peak hour (PHF). Merrimack's peak-day flow governs the day-scale elements: flow equalisation, storage, emergency lagoons and process buffering, where what matters is the worst day's volume, not the worst hour's rate. The biological process itself uses neither — it is sized on AAF plus loads (Section 12).")
para("So Merrimack is one of the STP's design checks, not the STP sizing method wholesale. The allocation above is standard practice; NWS's confirmation of the binding formula per element is on the kickoff list — the difference is roughly 40 percent of pipe capacity, which is not a rounding issue. The R0 model additionally applies a +20 percent weekly peak with no guideline source; it is flagged in Section 16 (and note that the R0 workbook's WW Generation series already contains this factor — Appendix B).")
para("Where do the Qm values come from? Qm is not an external data item: it is the computed mean (average) flow at the point being designed — the output of Step 5 accumulated through the network. Every pipe reach has its own Qm (the sum of all upstream zone flows), the STP inlet's Qm is the total Qadf, and each value comes from the flow model at that node. No measurement is needed to apply Peltier at design time; the STP inflow records requested from NWS serve to calibrate these computed values, not to replace them.")

# ================= 10 STEP 7 LOADS =================
H(1, "Step 7 — Organic (Pollution) Loads")
H(2, "What the loads are")
para("Flow says how much liquid arrives; loads say how much pollution it carries. The STP's biological reactors, aeration system, clarifiers and sludge line are sized by the daily mass of pollutants, in kilograms per day — not by concentration. Three parameters characterise domestic sewage:")
bullet("the mass of oxygen that bacteria consume in five days while degrading the biodegradable organic matter in the sample. It is the standard measure of 'how much food for bacteria' the sewage carries, and it directly sets the biological reactor volume and the oxygen (aeration energy) demand.", bold_lead="BOD5 (biochemical oxygen demand, 5-day): ")
bullet("the particulate matter suspended in the sewage. It settles in clarifiers and largely becomes primary/secondary sludge, so it sizes the clarifiers and the sludge handling line.", bold_lead="TSS (total suspended solids): ")
bullet("the oxygen equivalent of everything chemically oxidisable, measured in a fast chemical test. COD ≥ BOD5 always; their ratio measures biodegradability — domestic sewage runs COD/BOD = 1.8-2.2 (G203-p74), and a higher ratio warns of industrial or septic inputs that biology alone cannot remove.", bold_lead="COD (chemical oxygen demand): ")
H(2, "Why loads are fixed per person, not per litre")
para("A person excretes and washes away a nearly constant daily mass of organic matter regardless of how much water carries it. The guidelines therefore fix minimum per-capita loads (G203-p74): at least 60 g BOD5 and 80 g TSS per capita per day. Load calculation is then simply mass times population:")
equation(mr("L") + mr(" = ") + frac(mr("P × ") + ssub(mr("l"), mr("cap")), mr("1000")))
para("with L in kg/d, P the population and l(cap) the per-capita load in g/c/d. Concentration is always derived afterwards, never assumed:")
equation(mr("C") + mr(" = ") + frac(mr("L × 1000"), Q("adf")))
para("with C in mg/l. Because the mass is fixed while Oman's water use is low, the concentration comes out high: around 300-400 mg/l BOD5, versus about 200 mg/l in water-rich European systems. This matters in practice — a concentrated influent shifts the STP process choice (oxygen transfer, sludge yield, possible primary treatment) and makes the infiltration assumption visible in the influent quality data: excessive assumed infiltration would predict dilute sewage that the STP's laboratory records will contradict.")

# ================= 11 STEP 8 USES =================
H(1, "Step 8 — Which Flow Sizes What")
para("Each element of the system is governed by the flow that physically challenges it. A pipe never sees the annual average — it sees this morning's peak; a biological reactor barely notices the peak hour — bacteria respond to the daily mass. The table below assigns the governing flow to each element, and the notes below explain each line.")
table(["Design element", "Governing flow", "Additional rules"], [
 ["Gravity pipe sizing", "Peak (hourly) flow", "d/D ≤ 0.65 (D ≤ 350 mm) or ≤ 0.50 (D > 350 mm); v = 0.75-3.0 m/s; Table 11 minimum gradients (G203-p26-29)"],
 ["Pumping stations / force mains", "Peak flow, with any one pump out of service", "v = 1.0 (intermittent) to 2.5 m/s max (G203-p39, p50)"],
 ["STP hydraulic pass-through", "PHF", "screens, channels, distribution structures pass the peak hour (G203-p65-66)"],
 ["STP biological process", "AAF + organic loads", "reactor volume, aeration, clarifiers, sludge line (G203-p65-66)"],
 ["TE / TSE pipelines", "TSE production = 95% of STP inlet", "TE roughness penalty: epsilon +30%, C −10% vs potable values (G202-p104)"],
], widths=[1.9, 1.8, 2.8], font=8)
tab_caption("Governing flow for each design element")
bullet("Gravity pipes carry the peak flow only partly full: the depth-to-diameter limit (d/D) keeps an air space for ventilation and capacity reserve, the minimum velocity 0.75 m/s at peak keeps solids moving (self-cleansing), and the maximum 3.0 m/s prevents abrasion of the invert.")
bullet("Pump stations must deliver the peak with their largest unit out of service (duty/standby) — pumps fail, sewage does not stop.")
bullet("The STP splits in two: everything the sewage flows through is sized on PHF, while the biological process is sized on AAF plus the Step 7 loads. Sizing biology on PHF would roughly double the reactors for no benefit; sizing hydraulics on AAF would flood the works every morning.")
bullet("The TE network is designed to G202 criteria with the treated-effluent roughness penalty, because TSE carries residual solids and biofilm potential that potable water does not.")

# ================= 12 STEP 9 STP =================
H(1, "Step 9 — STP Incoming Flow, TSE and Sludge")
H(2, "The STP design flow")
para("The flow the treatment plant must be designed to receive is a defined sum (G203-p65 Table 29, G201-p73), expressed as:")
equation(Q("STP") + mr(" = 1.10 × (") + Q("adf") + mr(" + ") + Q("tank") + mr(")"))
para("read as: the network's average flow including its infiltration, plus the tankered deliveries arriving by road, all increased by a 10 percent operational allowance. The allowance covers what the estimate cannot see: metering error, unplanned connections, wet-year infiltration, and the operational reality that plants must work with one unit in maintenance. The hydraulic elements of the plant then apply the Step 6 peak factors on top of this average, while the biological elements combine it with the Step 7 loads.")
para("The resulting capacity places the plant in the G203-p65 size categories (Medium 500-20,000 m3/d, Large ≥ 20,000 m3/d), and — per the TOR — the Phase I capacity relative to 20,000 m3/d decides whether the new STP's preliminary design and EPC tender remain within this consultancy.")
H(2, "TSE — the product")
para("The plant returns 95 percent of its inflow as treated sewage effluent (G201-p73); the remaining 5 percent leaves with the sludge and process losses. TSE is not waste: it is the feedstock of the TE network that this project also designs, irrigating landscaping and farms. Its quality must meet Omani Class A standards where discharge to wadis is foreseen (G203-p73). The TSE volume from the equation above is therefore the input to the TE network's own demand-and-distribution design.")
H(2, "Sludge — the by-product")
para("Treatment concentrates the removed pollution into sludge — as a planning rate, about 0.25 kg of dry solids per m3 treated (R0). The sludge line thickens and dewaters it (volume reduction), stabilises it (odour and pathogen control), and directs it to its fate: beneficial reuse (soil conditioning, subject to quality) or landfill disposal. The TOR requires a sludge management strategy as a concept-stage deliverable, and the Step 7 TSS load is the primary input to its sizing.")
pagebreak()

# ================= 14 (Rev 3) COLEBROOK-WHITE PIPE HYDRAULICS =================
H(1, "How a Pipe Carries the Flow — Colebrook-White Hydraulics")
H(2, "Why this chapter exists")
para("Everything so far answers one question: how much sewage arrives. Steps 1-5 built the average flow, Step 6 peaked it, Step 7 loaded it. None of it tells you whether a DN250 laid at 4 mm/m can actually carry that peak. That is a hydraulics question, and this chapter closes the gap. It also matters practically: this is exactly the method the W4 network-design pipeline implements, pipe by pipe. The guideline accepts Colebrook-White or Manning (G203-p24); this project uses Colebrook-White throughout, with the fixed inputs the guideline prescribes.")
H(2, "The equation, piece by piece")
para("For a pipe running exactly full, Colebrook-White gives the velocity directly — no iteration, no friction-factor chart:")
equation(mr("V") + mr(" = −2 ") + sqrt(mr("2gDS")) + mr(" × ") + ssub(mr("log"), mr("10")) + mr("( ") +
         frac(ssub(mr("k"), mr("s")), mr("3.71 D")) + mr(" + ") +
         frac(mr("2.51 ν"), mr("D ") + sqrt(mr("2gDS"))) + mr(" )"))
para("In words: gravity pulls the water down the slope; the pipe wall and the water's own viscosity hold it back; V is where the two balance. The pieces:")
bullet("the full-bore velocity in m/s — what the pipe does when it runs exactly full.", bold_lead="V — ")
bullet("gravity, 9.81 m/s2.", bold_lead="g — ")
bullet("the internal pipe diameter in metres.", bold_lead="D — ")
bullet("the hydraulic gradient in m/m — for gravity sewers, simply the pipe slope. A Table 11 value in mm/m divided by 1000.", bold_lead="S — ")
bullet("the wall roughness, fixed at 1.5 mm for all pipe sizes and all materials (G203-p24, p28). This is deliberately not the catalogue smoothness of new PVC. A working sewer wall carries a grease-and-slime film, joints and small deposits, and after a few months a slimed PVC wall behaves much like a slimed concrete one. One value for everything is the guideline's honest way of saying so.", bold_lead="ks = 1.5 mm — ")
bullet("the kinematic viscosity of sewage, taken as water at 15 °C (G203-p25). Cooler water is thicker and flows slower, so this is the conservative design pick: real Omani sewage is warmer and will only flow faster than computed.", bold_lead="ν = 1.141 × 10−6 m2/s — ")
bullet("the left term is the wall-roughness resistance, the right term the viscous resistance. In sewer sizes the roughness term dominates; the viscous term only really matters in small pipes on flat slopes.", bold_lead="The two terms inside the log — ")
H(2, "Running partly full")
para("A gravity sewer is designed never to run full — so the full-bore V above is a reference point, not the operating condition. At a depth d in a pipe of diameter D, the water fills a circular segment. Everything follows from the angle θ that the water surface subtends at the pipe centre:")
equation(mr("θ") + mr(" = 2 ") + ssup(mr("cos"), mr("−1")) + mr("(1 − 2d/D)"))
equation(mr("A") + mr(" = ") + frac(ssup(mr("D"), mr("2")), mr("8")) + mr(" (θ − sin θ)") + mr(" ,    R = ") +
         frac(mr("A"), mr("P")) + mr(" = ") + frac(mr("D"), mr("4")) + mr(" (1 − ") + frac(mr("sin θ"), mr("θ")) + mr(")"))
para("A is the flow area, P = θD/2 the wetted perimeter, and R = A/P the hydraulic radius — flow area per unit of wall doing the braking. A full pipe has R = D/4, so replacing D with 4R in the Colebrook-White equation turns the full-bore formula into the any-depth formula. That substitution is the whole trick; there is no separate partial-flow theory.")
para("The behaviour that falls out is worth internalising (Figure 10). At half depth the velocity equals the full-bore velocity — half the area, but also half the wall, so R is unchanged and the flow does not notice. Velocity peaks about 14 percent above full-bore at d/D ≈ 0.81, capacity about 7 percent above full-bore at d/D ≈ 0.94; both then drop as the last of the air space disappears and the top wall starts braking too. So a sewer at three-quarters depth is not a sewer in trouble — it is a sewer near its hydraulic optimum.")
para("The guideline still caps the design depth at peak flow: d/D ≤ 0.65 for D ≤ 350 mm and d/D ≤ 0.50 for D > 350 mm (G203-p27 Tab 10). The air space is not waste — it ventilates the sewer (sulphide and odour control), absorbs the surge nobody predicted, and keeps capacity reserve precisely where the peak-factor estimate is shakiest. The stricter limit for the big pipes reflects how much more damage a surcharging trunk does.")
pic(CW_PARTIAL_CHART, w=5.9)
fig_caption("Partial-full behaviour of a circular sewer (Colebrook-White, DN315 at 4.0 mm/m): velocity and discharge relative to full-bore, with the d/D design limits and this chapter's sizing example marked")
H(2, "Where Table 11 actually comes from")
para("Every sewer designer carries the minimum-gradient table (G203-p29 Tab 11) around as if it were a separate rule. It is not. Table 11 is the Colebrook-White equation run backwards: for each diameter, find the slope at which the full-bore velocity equals the self-cleansing minimum of 0.75 m/s (G203-p26). Nothing else goes in. Here is DN200 at its tabulated 5.00 mm/m, step by step:")
bullet("2gDS = 2 × 9.81 × 0.200 × 0.00500 = 0.01962 m2/s2, so √(2gDS) = 0.1401 m/s.")
bullet("Roughness term: ks/(3.71 D) = 0.0015/0.742 = 0.00202.")
bullet("Viscous term: 2.51ν/(D·√(2gDS)) = 2.51 × 1.141e-6/(0.200 × 0.1401) = 1.02e-4 — twenty times smaller, as promised.")
bullet("Sum = 0.00212; log10 = −2.673.")
bullet("V = −2 × 0.1401 × (−2.673) = 0.749 m/s ≈ 0.75 m/s. That is the whole derivation.")
para("The same check on every row of Table 11 (computed live by this document's build script):")
table(["DN (mm)", "Table 11 minimum gradient (mm/m)", "CW full-bore velocity at that gradient (m/s)"],
      [[f"DN{dn}", f"{sm:.2f}", f"{cw_v(dn/1000.0, sm/1000.0):.3f}"] for dn, sm in TAB11],
      widths=[1.3, 2.6, 2.6], font=9)
tab_caption("G203 Table 11 minimum gradients reproduced by Colebrook-White at 0.75 m/s full-bore")
para("Every row lands between 0.74 and 0.77 m/s — the spread is only the rounding of the published gradients. Two consequences. First, if you trust the equation you can reproduce, extend or interpolate the table; the W4 pipeline does exactly this check on every pipe. Second, the table's logic is full-bore: at the actual (partial) peak depth the velocity differs, which is why the 0.75 m/s criterion is checked at peak flow in design, and why the guideline separately handles the near-empty case (below).")
pic(CW_TAB11_CHART, w=5.9)
fig_caption("Colebrook-White full-bore velocity at each Table 11 minimum gradient — the entire table reproduces the 0.75 m/s self-cleansing target")
H(2, "A worked sizing example")
para("A street collects Qpeak = 45 L/s (say, the Peltier-peaked outflow of a few zones) and the ground gives 4.0 m of fall per kilometre — S = 4.0 mm/m. The catalogue offers DN250 and DN315. In this tutorial DN is treated as the internal diameter; in detail design use the manufacturer's actual bore. Same equation, twice:")
bullet(f"√(2gDS) = √(2 × 9.81 × 0.250 × 0.004) = 0.1401 m/s; V = {cw_v(0.250, 0.004):.3f} m/s; A = 0.0491 m2; full-bore capacity Q = V × A = {cw_v(0.250, 0.004)*3.14159265*0.0625/4*1000:.1f} L/s. That is less than 45 L/s — the pipe cannot pass the peak even running completely full. Rejected before the depth check even starts.", bold_lead="DN250: ")
bullet(f"V = {cw_v(0.315, 0.004):.3f} m/s full-bore; A = 0.0779 m2; capacity {cw_v(0.315, 0.004)*3.14159265*0.315**2/4*1000:.1f} L/s. The 45 L/s peak uses 64 percent of that. Solving the partial-flow geometry for 45 L/s: the pipe runs at d/D ≈ 0.58 with V ≈ 0.96 m/s.", bold_lead="DN315: ")
para("Now the checks, all against Section 12's table: depth d/D = 0.58 ≤ 0.65 (DN315 is under the 350 mm threshold, G203-p27) — passes. Velocity 0.96 m/s sits between the 0.75 m/s self-cleansing minimum and the 3.0 m/s abrasion maximum (G203-p26-27) — passes, comfortably above the preferred 0.90 m/s too. The slope 4.0 mm/m exceeds DN315's Table 11 minimum of 2.70 mm/m — passes by construction. And at the d/D = 0.65 ceiling the pipe would carry 53 L/s, so there is 18 percent of growth headroom above the design peak. DN315 it is.")
table(["Check", "DN250 @ 4.0 mm/m", "DN315 @ 4.0 mm/m"], [
 ["Full-bore velocity (CW)", f"{cw_v(0.250, 0.004):.3f} m/s", f"{cw_v(0.315, 0.004):.3f} m/s"],
 ["Full-bore capacity", "38.1 L/s < 45 — fails", "70.4 L/s"],
 ["Depth at Qpeak = 45 L/s", "— (over-full)", "d/D ≈ 0.58 ≤ 0.65 ✓"],
 ["Velocity at Qpeak", "—", "≈ 0.96 m/s (0.75 ≤ v ≤ 3.0) ✓"],
 ["Capacity at d/D = 0.65", "—", "53 L/s → 18% headroom"],
 ["Verdict", "rejected", "selected"],
], widths=[2.2, 1.9, 2.4], font=9)
tab_caption("Worked pipe-sizing example: Qpeak = 45 L/s on a 4.0 mm/m available fall")
H(2, "At the network heads: the tractive-force minimum")
para("One complement, briefly. At the top ends of the network the peak flow is 1-2 L/s and 'self-cleansing at 0.75 m/s full-bore' loses its meaning — the pipe runs a few centimetres deep and Table 11's flattest gradients will not drag sand along the invert at such depths. For these reaches the guideline requires the steeper of the self-cleansing and tractive-force minimum gradients (G203-p27 §4.2.2):")
equation(ssub(mr("S"), mr("min")) + mr(" = 2.33 × ") + ssup(mr("10"), mr("−4")) + mr(" × ") +
         ssup(mr("τ"), mr("1.23")) + mr(" × ") + ssup(mr("Q"), mr("−0.461")))
para("with Q in m3/s and τ the tractive tension in Pa — the shear stress the trickle must exert on the invert to move grit. G203 gives no numeric design value for τ; this project carries τ = 1 Pa as a tagged pending assumption (GAP-9), literature-based, to be confirmed with NWS. At Q = 1.0 L/s the formula gives Smin = 5.6 mm/m — steeper than the flattest use of Table 11's DN200 row (5.00 mm/m), which is precisely the point: the smaller the flow, the more slope it needs, and the table alone will not tell you that. The W4 pipeline applies this check automatically at every network head; the full treatment lives in the W4 design criteria, not here.")
pagebreak()

# ================= 15 WORKED EXAMPLE =================
H(1, "Worked Example")
para("The example is constructed for this tutorial (it is not taken from a source document): a hypothetical settlement of 10,000 persons, served by 25 km of new sewers, lying inland with deep groundwater — deliberately round numbers so each step is easy to follow. Values are rounded for readability.", italic=True)
para("Walk-through of the solution:")
bullet("Steps 1-2: the population is given (10,000). Water demand = 164 l/c/d domestic × 1.36 (adding 22% + 14%) = 223 l/c/d, i.e. 2,230 m3/d of water supplied.")
bullet("Step 3: sewage per person = 164 × 0.85 + 59.1 × 0.54 = 171.3 l/c/d, i.e. 1,713 m3/d entering the sewers.")
bullet("Step 4: infiltration = 720 L/d/km × 25 km = 18 m3/d — barely 1 percent of the sewage, as expected for a new inland network. (Under the R0's 10 percent rule it would be 171 m3/d — the difference the reconciliation register tracks.) No tankers: the settlement is fully networked.")
bullet("Step 5: Qadf = 1,713 + 18 = 1,731 m3/d. Divide by 86.4 to convert m3/d to L/s: 20.0 L/s.")
bullet("Step 6: Peltier PF = 1.5 + 1/sqrt(20.0) = 1.72, so peak-hour flow = 1.72 × 20.0 = 34.5 L/s — this sizes the outfall sewer. Merrimack: Qadf = 1.731 Ml/d, so Qpdf = 2.65 × 1.731^0.879 = 4.29 Ml/d (49.7 L/s, an implied factor of 2.48) — the peak-day cross-check.")
bullet("Step 7: loads = 10,000 × 60 / 80 / 120 g/d = 600 kg BOD5, 800 kg TSS, 1,200 kg COD per day (COD taken at 2.0 × BOD). Concentrations = load × 1000 / 1,731: BOD5 347 mg/l, TSS 462 mg/l, COD 693 mg/l — correctly in the concentrated range expected for Oman.")
bullet("Steps 8-9: STP design flow = 1.10 × 1,731 = 1,904 m3/d (Medium category); TSE production = 0.95 × 1,904 = 1,809 m3/d available to the TE network; sludge ≈ 0.25 × 1,904 = 476 kg/d dry solids.")
table(["#", "Step", "Calculation", "Result"], [
 ["1", "Population", "given", "10,000 cap"],
 ["2", "Water demand", "164 × 1.36", "223 l/c/d → 2,230 m3/d"],
 ["3", "Wastewater return", "164 × 0.85 + 59.1 × 0.54", "171.3 l/c/d → 1,713 m3/d"],
 ["4", "Infiltration (new network)", "720 L/d/km × 25 km", "+18 m3/d (≈ 1%)"],
 ["5", "Average flow Qadf", "1,713 + 18", "1,731 m3/d = 20.0 L/s"],
 ["6a", "Peltier peak (hourly)", "PF = 1.5 + 1/√20.0 = 1.72", "peak 34.5 L/s"],
 ["6b", "Merrimack peak (daily)", "2.65 × 1.731^0.879", "4.29 Ml/d = 49.7 L/s (PF 2.48)"],
 ["7", "Loads BOD5 / TSS / COD", "10,000 × 60 / 80 / 120 g/d", "600 / 800 / 1,200 kg/d"],
 ["7b", "Concentrations", "load × 1000 / Qadf", "BOD5 347, TSS 462, COD 693 mg/l"],
 ["8", "STP design flow", "1.10 × 1,731", "1,904 m3/d"],
 ["9", "TSE production", "0.95 × 1,904", "1,809 m3/d"],
], widths=[0.45, 1.9, 2.5, 1.9], font=9)
tab_caption("Worked example — hypothetical 10,000-person settlement (illustrative)")
H(2, "Sanity checks — and why they work")
para("Three quick checks catch most calculation errors, each grounded in the physics of the previous sections:")
bullet("Because the per-capita mass is fixed (Step 7) while Omani water use is low, dividing the two must land in this range. A dilute result (~200 mg/l) almost always means the flow is overstated — typically an inflated infiltration or per-capita water figure.", bold_lead="BOD5 concentration 300-400 mg/l. ")
bullet("The Peltier structure guarantees PF > 1.5, and the guideline caps it at 5.0. A computed PF outside this band means the flow unit is wrong (m3/d fed into a formula expecting L/s is the classic error).", bold_lead="Peak factor between 1.5 and 5.0. ")
bullet("At 720 L/d/km, even hundreds of kilometres of sewer add only a few percent. If infiltration rivals the sewage flow in a new inland design, an existing-network rate (10-40%) has leaked into the wrong context.", bold_lead="Infiltration a small fraction of sewage for a new inland network. ")

# ================= 14 RECONCILIATION =================
H(1, "Guideline vs Inception R0 — Reconciliation Register")
para("The R0 demand model and the guidelines agree on the core parameters; the differences below are carried openly until confirmed with NWS at kickoff. None of them invalidates the method of this tutorial — they change input values, and the tutorial states at each point which value it uses.")
table(["Item", "Guideline (binding)", "R0 adopted", "Position"], [
 ["Domestic per-capita demand", "164 l/c/d (G201-p60)", "163.5 computed from billing", "Consistent — use 164"],
 ["Return ratios", "85% / 54% (G201-p71)", "Same", "Consistent"],
 ["Infiltration", "720 L/d/km, new networks (G201-p72)", "10% of wastewater flow", "Guideline value as design basis; 10% kept as conservative upper sensitivity — NWS to confirm"],
 ["Peaking", "Peltier / Merrimack, PF ≤ 5 (G201-p71-72)", "+20% weekly peak added", "No guideline source — basis to be confirmed"],
 ["Tanker catchment", "Not specified", "25 km radius", "R0 assumption — do not mix into design until NWS confirms; site visit shows up to 150 km actual"],
 ["STP margin / TSE ratio", "+10% / 95% (G201-p73)", "Same", "Consistent"],
 ["Sludge rate", "Not specified", "0.25 kg/m3", "R0 planning value — refine at process design"],
], widths=[1.5, 2.0, 1.6, 1.9], font=8)
tab_caption("Reconciliation of guideline values against the R0 demand model")

# ================= 15 CONCLUSION =================
H(1, "Conclusions")
bullet("The sewage flow and load calculation is a fully determined nine-step chain: every parameter traces to a cited guideline value, an NCSI statistic, or a flagged assumption — nothing is discretionary once the inputs are fixed.")
bullet("Population is the root input and its two routes serve different decisions: census projection for the phased facilities (STP), plot saturation for the once-only network. Their consistency (projection below ceiling, zone by zone) is itself a design check.")
bullet("Flows and loads are independent tracks that meet at the STP: flows follow the water, loads follow the people. Oman's low water use therefore produces concentrated sewage (300-400 mg/l BOD5), which the process design must expect.")
bullet("The design-governing quantities are: Route B saturation flow for pipes; Route A dated flows for STP phases; PHF for everything hydraulic; AAF + loads for everything biological; and the +10% allowance at the plant gate.")
bullet("Four input questions remain open — infiltration basis, peaking basis, tanker catchment, occupancy rate — and all four are data/policy questions for NWS, not methodological gaps.")
H(1, "Recommendations")
bullet("Request at kickoff: NCSI housing-unit counts for the occupancy rate (not in the R0 package), electricity account counts per settlement (for sub-settlement disaggregation), and NWS confirmation of the infiltration and peaking bases.")
bullet("Adopt the guideline infiltration (720 L/d/km) as the design basis for the new network and carry the R0's 10 percent as an upper sensitivity, so that STP phasing is not driven by phantom flow.")
bullet("Use Peltier for hydraulic sizing (current IMP method) with Merrimack as a peak-day cross-check, pending NWS confirmation.")
bullet("Obtain an NWS policy decision on the tanker catchment obligation (25 km assumption vs the observed 150 km deliveries) before sizing the tanker reception facility.")
bullet("Validate the aggregate non-domestic allowance (22% + 14%) against the spatially-assigned Table 12 loads once the land-use model is built, and investigate any large discrepancy.")

# ================= 16 REFERENCES =================
H(1, "References")
bullet("PAM-GUD-201 — General Design Guidelines v1.0, Rev 01, Nama Water Services (cited G201-p##).")
bullet("PAM-GUD-202 — Potable Water & Treated Sewerage Effluent (TSE) Design Guidelines v1.0, Rev 01, Nama Water Services (cited G202-p##).")
bullet("PAM-GUD-203 — Wastewater Design Guidelines v1.0, Rev 01, Nama Water Services (cited G203-p##).")
bullet("Terms of Reference, Section 03, Tender T/2719110/2025, Nama Water Services (cited TOR).")
bullet("Ibri Sewer Demand R0 workbook and Inception Report R0, Renardet, August 2026 (cited R0).")
bullet("NCSI population and housing statistics, National Centre for Statistics and Information, data portal (as transmitted in the R0 workbook).")

# ================= APPENDIX =================
H(1, "Appendix A — NCSI Population Series, Ibri Wilayat")
yrs = sorted(NCSI)
table(["Year"] + [str(y) for y in yrs],
      [["Population"] + [f"{NCSI[y]:,}" for y in yrs]], font=7)
tab_caption("Full NCSI population series for Ibri wilayat as carried in the R0 workbook (2021-2033)")
para("Growth rates implied by the series decline from 2.97 %/yr (2022) to 1.29 %/yr (2033). The connected-population ratio in the R0 workbook (share of population on the potable network) rises from 46.7% (2021) to 74.6% (2033).")

# ================= APPENDIX B =================
H(1, "Appendix B — The R0 Demand Workbook, Tab by Tab")
para("This appendix documents the structure, data and equations of the R0 demand workbook ('Ibri Sewer Demand R0 2026 08 03.xlsx'), so that its results can be reproduced and audited. The description below was established by tracing the workbook's cell formulas directly; the master equation of B.6 is the exact formula implemented per settlement and year.")
H(2, "B.1  'Project Pop Settlements' — the population engine")
para("One row per project settlement (50 settlements), one column per year from 2023 to 2100. Each cell looks up the settlement's projected population from the wilayat-level settlement projection ('Ibri Pop Settlements', 240 settlements), which in turn distributes the NCSI wilayat series (Appendix A) over settlements by their constant census shares — exactly Route A of Section 5. Key values: project total 116 thousand (2024), 238 thousand (2055), 691 thousand (2100); largest settlements IBRI town, AD DARIZ, AL ARAQI (Figure 5). The projection applies growth throughout — it contains no saturation ceiling (see Section 5.4).")
H(2, "B.2  'Project W Pop Connected' — who is on the water network")
para("Same layout; each cell multiplies the settlement population by the wilayat-wide connected-population ratio c(y) (tab 'Pop_Wilayat', row 'W_Pop_Connected'):")
equation(ssub(mr("P"), mr("conn,s")) + mr("(y) = ") + ssub(mr("P"), mr("s")) + mr("(y) × c(y)"))
para("The ratio ramps linearly from 46.7% (2021) through 74.6% (2033), reaching 100% in 2070 and staying there — i.e. the model assumes universal water-network coverage only from 2070 onward. Only the connected population generates billed water demand in the model; the remainder appears through the tanker stream.")
pic(RAMP_CHART, w=5.9)
fig_caption("Connected-population ratio c(y) in the R0 model: 46.7% (2021) ramping to 100% (2070)")
H(2, "B.3  'Project Tankers' — the unconnected stream")
para("One row per settlement, columns 2024-2100, values in m3/d of tankered water-consumption equivalent. Two modelling choices to note: (i) only settlements within 25 km of the STP by road contribute (criteria tab, 'Maximum Distance from STP by road' = 25 km) — an R0 assumption with no guideline source, flagged in Section 16; (ii) the series is constant at its 2024 estimate for every year — tanker volumes are not grown nor phased out, which is conservative for the STP but inconsistent with the guideline's expectation that coverage reaches 100% and tankering declines (G201-p73). Both points are on the kickoff confirmation list.")
pic(TANKER_CHART, w=5.9)
fig_caption("Ten largest tanker-served settlements in the R0 model (2024 volumes, project total 942 m3/d)")
H(2, "B.4  'Water Demand Criteria' — the demand rate library")
para("Holds, per governorate and year: the domestic per-capita consumption computed from billing (Adh Dhahirah: 163.5 l/c/d from 2024 onward, the value used throughout the model), the governmental consumption ratio (Adh Dhahirah: 0.140, i.e. 14.0% of domestic) and the non-domestic ratio (Adh Dhahirah: 0.218, i.e. 21.8% of domestic). These are the workbook's counterparts of the guideline values 164 l/c/d, 14% and 22% (G201-p60-61) — agreement is essentially exact, because both derive from the same IMP/billing basis.")
pic(LPCD_CHART, w=5.9)
fig_caption("Measured domestic consumption by governorate, 2024 (Adh Dhahirah highlighted; workbook 'Water Demand Criteria')")
H(2, "B.5  'Wastewater Criteria' — the conversion constants")
table(["Constant", "Value", "Role"], [
 ["Domestic return ratio", "0.85", "share of domestic water becoming sewage (G201-p71)"],
 ["Tanker return ratio", "0.85", "same, for tankered consumption"],
 ["Non-domestic return ratio", "0.54", "share of non-domestic water becoming sewage"],
 ["Governmental return ratio", "0.54", "share of governmental water becoming sewage"],
 ["Infiltration ratio", "0.10", "added to sewage flow; workbook note: applied only to settlements with an existing network"],
 ["Max tanker distance", "25 km", "tanker catchment cut-off (B.3)"],
 ["Weekly peak", "0.20", "+20% factor applied inside the WW Generation series (B.6)"],
 ["STP margin coefficient", "0.10", "operational allowance, applied at STP sizing"],
 ["TE production rate", "0.95", "TSE = 95% of STP inlet"],
 ["Sludge rate", "0.25 kg/m3", "planning sludge production"],
], widths=[2.0, 1.0, 3.5], font=8)
tab_caption("Constants of the 'Wastewater Criteria' tab (all marked 'can be adjusted' in the workbook)")
H(2, "B.6  'Project WW Generation' — the master equation")
para("Each cell of the WW Generation tab computes, for settlement s and year y (traced from the cell formulas):")
equation(ssub(mr("WWG"), mr("s")) + mr(" = ") + mr("[ ") + ssub(mr("P"), mr("conn,s")) + mr(" × ") + frac(mr("L"), mr("1000")) + mr(" × (") + ssub(mr("R"), mr("dom")) + mr(" + ") + ssub(mr("R"), mr("nd")) + mr("×") + ssub(mr("r"), mr("nd")) + mr(" + ") + ssub(mr("R"), mr("gov")) + mr("×") + ssub(mr("r"), mr("gov")) + mr(") × (1 + i) + ") + ssub(mr("T"), mr("s")) + mr(" × ") + ssub(mr("R"), mr("tank")) + mr(" ] × (1 + w)"))
table(["Symbol", "Meaning", "Value"], [
 ["P(conn,s)", "connected population of settlement s (B.2)", "per settlement/year"],
 ["L", "domestic per-capita demand, l/c/d (B.4)", "163.5"],
 ["R(dom), R(tank)", "domestic / tanker return ratios", "0.85"],
 ["R(nd), R(gov)", "non-domestic / governmental return ratios", "0.54"],
 ["r(nd), r(gov)", "non-domestic / governmental consumption ratios (B.4)", "0.218 / 0.140"],
 ["i", "infiltration ratio (network settlements)", "0.10"],
 ["T(s)", "tankered volume of settlement s, m3/d (B.3)", "constant 2024 values"],
 ["w", "weekly peak factor", "0.20"],
], widths=[1.3, 3.5, 1.7], font=8)
tab_caption("Symbols of the R0 master equation")
para("Two audit findings matter when using these numbers:")
bullet("the (1 + w) factor means every WWG value is 20 percent above the average-flow chain of this tutorial. R0's WWG is therefore closer to a 'design weekly flow' than to Qadf, and it must not be fed into formulas expecting an average (e.g. Peltier's Qm) without first removing the factor.", bold_lead="The WW Generation series already includes the +20% weekly peak: ")
bullet("in this workbook infiltration enters as a ratio of sewage flow (B.5), not as the guideline's 720 L/d/km of pipe — the deviation discussed in Sections 8 and 16.", bold_lead="Infiltration is inside the generation formula: ")
pic(WWG_CHART, w=6.0)
fig_caption("R0 wastewater generation to 2100 — three largest settlements and project total ('Project WW Generation'; values include the +20% weekly peak)")
para("With these two caveats, the workbook chain is exactly the chain of this tutorial: Steps 1 (B.1-B.2), 2 (B.4), 3-4 (B.5, B.3) and 5 (B.6), with the peaking and STP steps applied downstream of it. The project WW generation reaches about 14,200 m3/d in 2024, 46,500 m3/d by 2055 and 157,000 m3/d by 2100 (weekly-peak basis; divide by 1.2 for the average-flow equivalent).")

# ---- populate List of Figures / List of Tables at the front markers ----
def fill_list(marker, items):
    for it in items:
        p = doc.add_paragraph(it)
        p.paragraph_format.space_after = Pt(2)
        for r in p.runs: r.font.size = Pt(10)
        marker._p.addprevious(p._p)
fill_list(LOF_MARK, FIGS)
fill_list(LOT_MARK, TABS)

doc.save(OUT)
print("saved", OUT, "| figures:", len(FIGS), "| tables:", len(TABS))
