# -*- coding: utf-8 -*-
"""W8 methodology report -> Word (.docx), rendered from the shared content blocks.

INTERNAL working document: clean professional styling, NOT the client Sample.docx
letterhead (rule 5 applies to client-facing reports; this one is for the design team).
Content and every number come from report_content.build() — the same source the PDF
renderer uses — so the two formats cannot drift.

Re-run: python W8/report/make_methodology_docx.py
(PDF: make_methodology_pdf.py — reportlab, no Word dependency.)
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import report_content as RC

OUT = os.environ.get("W8_DOCX_OUT") or os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "W8_Sewer_Network_Design.docx")

doc = Document()

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
for lvl, sz, col in ((1, 15, "1F3864"), (2, 12.5, "2E5A88")):
    h = doc.styles[f"Heading {lvl}"]
    h.font.name = "Calibri"
    h.font.size = Pt(sz)
    h.font.color.rgb = RGBColor.from_string(col)
    h.font.bold = True
for s_ in doc.sections:
    s_.top_margin = s_.bottom_margin = Inches(0.8)
    s_.left_margin = s_.right_margin = Inches(0.85)


def para(text="", align=None, bold=None, size=None, color=None, space_after=6, italic=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)


def table(headers, rows, widths=None, font=8.6):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(font)
        shade(c, "D9E2F3")
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(font)
    if widths:
        for i, w in enumerate(widths):
            for r_ in t.rows:
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def caption(text):
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(text)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("555555")


def toc():
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click > Update Field to build the Table of Contents."
    r.append(t)
    fld.append(r)
    p._p.append(fld)


for b in RC.build():
    kind = b[0]
    if kind == "cover":
        c = b[1]
        para("", space_after=90)
        para(c["eyebrow"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13, color="555555")
        para(c["title"], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=25, color="1F3864",
             space_after=8)
        para(c["subtitle"], align=WD_ALIGN_PARAGRAPH.CENTER, size=15, color="2E5A88",
             space_after=40)
        para(c["note"], align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color="777777")
        para(c["date"], align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color="777777", space_after=36)
        table(["Item", "Value"], c["facts"], widths=[2.2, 4.0], font=10)
    elif kind == "toc":
        toc()
    elif kind == "pagebreak":
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    elif kind == "h1":
        doc.add_paragraph(b[1], style="Heading 1")
    elif kind == "h2":
        doc.add_paragraph(b[1], style="Heading 2")
    elif kind == "p":
        para(b[1])
    elif kind == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b[1]).bold = True
        p.add_run(b[2])
        p.paragraph_format.space_after = Pt(4)
    elif kind == "table":
        table(b[1], b[2], widths=b[3])
    elif kind == "img":
        path, want, cap = b[1], b[2], b[3]
        if os.path.exists(path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(path, width=Inches(want))
            caption(cap)
    elif kind == "img2":
        p1, p2, cap = b[1], b[2], b[3]
        if os.path.exists(p1) and os.path.exists(p2):
            t = doc.add_table(rows=1, cols=2)
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for cell, path, w in ((t.rows[0].cells[0], p1, 2.5), (t.rows[0].cells[1], p2, 2.9)):
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.paragraphs[0].add_run().add_picture(path, width=Inches(w))
            caption(cap)

doc.save(OUT)
print("saved", OUT)
