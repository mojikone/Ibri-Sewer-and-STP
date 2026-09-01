"""Document furniture for T03 - styles, headings, tables, captions, TOC.

Everything writes in document flow (before the trailing sectPr), so content
appears where it is called rather than collecting at the end of the file.
"""
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BLUE = RGBColor(0x1F, 0x3B, 0x63)
MID = RGBColor(0x2E, 0x62, 0x9E)
GREY = RGBColor(0x5A, 0x5A, 0x5A)
RED = RGBColor(0xA6, 0x1B, 0x1B)

_counters = {"fig": 0, "tab": 0, "eq": 0}


# ------------------------------------------------------------------ document
def new_document():
    d = Document()
    s = d.sections[0]
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.left_margin = s.right_margin = Cm(2.2)
    s.top_margin = Cm(2.2)
    s.bottom_margin = Cm(2.0)

    n = d.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12

    sect = s._sectPr
    sect.append(parse_xml(
        '<w:footnotePr xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">'
        '<w:numRestart w:val="eachPage"/></w:footnotePr>'))

    for name, size, colour, before in (
        ("Heading 1", 17, BLUE, 18),
        ("Heading 2", 13.5, MID, 14),
        ("Heading 3", 11.5, MID, 10),
    ):
        st = d.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True
    return d


def _flow(d, element):
    body = d.element.body
    sect = body.find(qn("w:sectPr"))
    if sect is not None:
        sect.addprevious(element)
    else:
        body.append(element)
    return element


# -------------------------------------------------------------------- blocks
def _outline(par, lvl):
    """Set the contents level without changing how the heading looks."""
    ppr = par._p.get_or_add_pPr()
    for old in ppr.findall(qn("w:outlineLvl")):
        ppr.remove(old)
    e = OxmlElement("w:outlineLvl")
    e.set(qn("w:val"), str(lvl))
    ppr.append(e)
    return par


def h(d, level, text, page_break=False):
    """A numbered section heading. Sections sit one level below Parts."""
    if page_break:
        pagebreak(d)
    _step["n"] = 0          # every heading starts a fresh list context
    par = d.add_heading(text, level)
    _outline(par, level)    # Heading 1 -> contents level 2, and so on
    return par


def part(d, letter, title):
    """A part divider: its own page, and the top level of the contents."""
    pagebreak(d)
    for _ in range(6):
        p(d, "", space_after=0)
    rule = p(d, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    _rule(rule)
    lab = p(d, "PART " + letter, align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True, size=15, colour=MID, space_after=10)
    ttl = d.add_paragraph()
    ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ttl.paragraph_format.space_after = Pt(6)
    r = ttl.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = BLUE
    rule2 = p(d, "", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _rule(rule2)
    _outline(ttl, 0)        # the part is the top level of the contents
    _step["n"] = 0
    pagebreak(d)            # the divider takes a page of its own
    return ttl


def _rule(par):
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    e = OxmlElement("w:bottom")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), "8")
    e.set(qn("w:space"), "4")
    e.set(qn("w:color"), "2E629E")
    pbdr.append(e)
    ppr.append(pbdr)
    return par


def p(d, text="", bold=False, italic=False, size=None, colour=None,
      align=None, space_after=None, style=None):
    par = d.add_paragraph(style=style)
    if align is not None:
        par.alignment = align
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    if text:
        r = par.add_run(text)
        r.bold, r.italic = bold, italic
        if size:
            r.font.size = Pt(size)
        if colour:
            r.font.color.rgb = colour
    return par


def rich(d, *parts, align=None, space_after=None):
    """rich(d, ("plain ", {}), ("bold", {"bold": True}), ...)"""
    par = d.add_paragraph()
    if align is not None:
        par.alignment = align
    if space_after is not None:
        par.paragraph_format.space_after = Pt(space_after)
    for text, fmt in parts:
        r = par.add_run(text)
        r.bold = fmt.get("bold", False)
        r.italic = fmt.get("italic", False)
        if fmt.get("size"):
            r.font.size = Pt(fmt["size"])
        if fmt.get("colour"):
            r.font.color.rgb = fmt["colour"]
        if fmt.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(fmt.get("size", 9.5))
    return par


def bullet(d, text, lead=None, level=0):
    par = d.add_paragraph(style="List Bullet")
    par.paragraph_format.left_indent = Cm(0.6 + 0.5 * level)
    par.paragraph_format.space_after = Pt(3)
    if lead:
        par.add_run(lead).bold = True
    par.add_run(text)
    return par


_step = {"n": 0}


def numbered(d, text, lead=None, restart=False):
    """Manually numbered step. Word's List Number style continues numbering
    across the whole document, so the number is written explicitly instead."""
    if restart:
        _step["n"] = 0
    _step["n"] += 1
    par = d.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.9)
    par.paragraph_format.first_line_indent = Cm(-0.9)
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(f"{_step['n']}.   ")
    r.bold = True
    if lead:
        par.add_run(lead).bold = True
    par.add_run(text)
    return par


def pagebreak(d):
    d.add_page_break()


def shade(par, hexfill):
    par._p.get_or_add_pPr().append(
        parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/'
                  f'wordprocessingml/2006/main" w:val="clear" w:color="auto" '
                  f'w:fill="{hexfill}"/>'))


def callout(d, title, text, fill="FFF4E5", colour=RED):
    """A boxed note - used for guideline defects and cautions."""
    par = d.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(8)
    par.paragraph_format.left_indent = Cm(0.3)
    r = par.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = colour
    r.font.size = Pt(10)
    par.add_run(text).font.size = Pt(10)
    shade(par, fill)
    _box(par)
    return par


def _box(par):
    ppr = par._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "6")
        e.set(qn("w:space"), "6")
        e.set(qn("w:color"), "D9A441")
        pbdr.append(e)
    ppr.append(pbdr)


# -------------------------------------------------------------------- tables
def table(d, headers, rows, widths=None, font=9, header_fill="1F3B63",
          align_right=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    align_right = align_right or set()

    for i, htxt in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        par = cell.paragraphs[0]
        par.paragraph_format.space_after = Pt(2)
        r = par.add_run(str(htxt))
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell._tc.get_or_add_tcPr().append(parse_xml(
            f'<w:shd xmlns:w="http://schemas.openxmlformats.org/'
            f'wordprocessingml/2006/main" w:val="clear" w:color="auto" '
            f'w:fill="{header_fill}"/>'))

    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            par = cells[i].paragraphs[0]
            par.paragraph_format.space_after = Pt(2)
            if i in align_right:
                par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            txt = "" if val is None else str(val)
            # **segments** render bold, wherever they sit in the cell
            for k, seg in enumerate(txt.split("**")):
                if not seg:
                    continue
                r = par.add_run(seg)
                r.font.size = Pt(font)
                r.bold = (k % 2 == 1)

    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)

    # repeat the header on every page, and never split a row across a page:
    # a broken row orphans its label and the table stops being readable
    hdr = t.rows[0]._tr.get_or_add_trPr()
    hdr.append(parse_xml('<w:tblHeader xmlns:w="http://schemas.openxmlformats.'
                         'org/wordprocessingml/2006/main"/>'))
    for row in t.rows:
        trpr = row._tr.get_or_add_trPr()
        trpr.append(parse_xml('<w:cantSplit xmlns:w="http://schemas.'
                              'openxmlformats.org/wordprocessingml/2006/main"/>'))
    return t


# ------------------------------------------------------------------ captions
def fig_caption(d, text):
    _counters["fig"] += 1
    par = d.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_after = Pt(10)
    r = par.add_run(f"Figure {_counters['fig']}   {text}")
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = GREY
    return _counters["fig"]


def tab_caption(d, text):
    _counters["tab"] += 1
    par = d.add_paragraph()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(3)
    r = par.add_run(f"Table {_counters['tab']}   {text}")
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = MID
    return _counters["tab"]


def next_eq():
    _counters["eq"] += 1
    return str(_counters["eq"])


def chart(d, name, width_cm=13.0):
    """Place a chart from img/ by name. Missing charts are skipped rather than
    breaking the build, so the report can be produced while charts are added."""
    import os
    path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "img", name + ".png")
    if not os.path.exists(path):
        return None
    return picture(d, path, width_cm)


def picture(d, path, width_cm=16.0):
    par = d.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(2)
    par.add_run().add_picture(path, width=Cm(width_cm))
    return par


# ---------------------------------------------------------------------- misc
def toc(d, levels="1-3"):
    par = d.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f'TOC \\o "{levels}" \\h \\z \\u')
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose Update Field to build the contents."
    run.append(t)
    fld.append(run)
    par._p.append(fld)
    return par


def footer_pagenum(d, left_text):
    for section in d.sections:
        ftr = section.footer.paragraphs[0]
        ftr.text = ""
        ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ftr.add_run(left_text + "     ")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        for instr in ("PAGE",):
            fld = OxmlElement("w:fldSimple")
            fld.set(qn("w:instr"), instr)
            ftr._p.append(fld)
