# -*- coding: utf-8 -*-
"""R1 — Concept Screening Report styled on Data/sample report/Sample.docx (python-docx).
Client-facing. Criteria refs: PAM-GUD-203 (p##) / PAM-GUD-201 (G1-p##)."""
import csv, os, copy, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SAMPLE = r"D:\Mojtaba\Renardet\2621 Ibri Sewer STP\Data\sample report\Sample.docx"
W2 = r"D:\Mojtaba\Renardet\2621 Ibri Sewer STP\Hydraulic\Claude\W2"
OUT = os.path.join(W2, "report", "Ibri_Concept_Screening_R1.docx")
MAPS = os.path.join(W2, "img", "maps")
PROJ = "Consultancy Services for Design and Supervision for STP, Sewer & TE Networks Systems in Ibri"
SUB = "Concept Screening Report"

doc = Document(SAMPLE)

# ---- wipe body (keep final sectPr) ----
body = doc.element.body
sect = body.find(qn("w:sectPr"))
for el in list(body):
    if el is not sect:
        body.remove(el)

# ---- headers/footers: swap project strings ----
def patch_part(part):
    xml = part._element
    first = True
    for t in xml.iter(qn("w:t")):
        if first:
            t.text = PROJ + " — " + SUB
            first = False
        else:
            t.text = ""
for s in doc.sections:
    for hd in (s.header, s.first_page_header, s.even_page_header):
        if hd is not None: patch_part(hd)

# ---- helpers ----
def para(text, style=None, align=None, bold=None, size=None, color=None, space_after=6):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    if bold is not None: r.bold = bold
    if size: r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p
def H(level, text):
    return doc.add_paragraph(text, style=f"Heading {level}")
_styles = {s_.name for s_ in doc.styles}
def bullet(text):
    if "List Bullet" in _styles:
        return doc.add_paragraph(text, style="List Bullet")
    p = doc.add_paragraph("• " + text)
    p.paragraph_format.left_indent = Inches(0.3)
    return p
def caption(text):
    try:
        p = doc.add_paragraph(text, style="Caption")
    except KeyError:
        p = para(text, bold=True, size=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p
def pic(path, w=6.4):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(w))
def shade(cell, hexc):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd"); sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc)
    tcPr.append(sh)
def table(headers, rows, widths=None, font=9):
    t = doc.add_table(rows=1, cols=len(headers))
    try: t.style = "Table Grid"
    except KeyError: pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(font)
        shade(c, "D9E2F3")
    for row in rows:
        cs = t.add_row().cells
        for i, v in enumerate(row):
            cs[i].text = ""
            r = cs[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(font)
    if widths:
        for i, w in enumerate(widths):
            for r_ in t.rows:
                r_.cells[i].width = Inches(w)
    return t
def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def toc():
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "Right-click and Update Field to refresh the Table of Contents."
    r.append(t); fld.append(r); p._p.append(fld)

# ================= COVER =================
para("", space_after=60)
para("Sultanate of Oman", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16)
para("Nama Water Services Company SAOC (NWS / OWWSC)", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14)
para("", space_after=40)
try:
    para(PROJ, style="Title", align=WD_ALIGN_PARAGRAPH.CENTER)
except KeyError:
    para(PROJ, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24)
para("Tender No. T/2719110/2025", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("", space_after=40)
para(SUB, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=22, color="0000CC")
para("Revision R1 — July 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
para("", space_after=80)
para("Renardet S.A. & Partners Consulting Engineers", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
para("Project No. 2621", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
pagebreak()
para("Table of Contents", bold=True, size=16, color="0000CC")
toc()
pagebreak()

# ================= 1 EXECUTIVE SUMMARY =================
H(1, "Executive Summary")
para("This report presents a desk-top concept screening of the wastewater collection strategy for Ibri Wilayat, prepared ahead of the project kick-off. Its purpose is to establish, on the basis of the available terrain model, cadastral and road data, and strictly in accordance with the NWS design guidelines (PAM-GUD-203 and PAM-GUD-201), the broad shape of the future sewer system: where the flows are generated, how they can reach the existing Ibri STP site, which areas can be served by gravity, where sewage lifting stations (SLS) are unavoidable, and what the ultimate flows imply for the treatment capacity.")
para("The principal findings are:")
bullet("Ibri is fundamentally a gravity catchment. The main agglomeration lies 45 to 90 m above the existing STP site, providing a median available fall of about 3.6 m/km along the natural drainage direction — several times the minimum gradient required for large trunk sewers (0.75 mm/m for DN900 and larger, Table 11 of PAM-GUD-203). A gravity trunk system of approximately 22 km of main trunk and 172 km of trunk branches can serve the entire boundary.")
bullet("The serviceable land use comprises 53,503 plots (43,722 residential). Applying the PAM-GUD-201 methodology (164 l/c/d for Adh Dhahirah, non-domestic and governmental ratios, return rates, linear infiltration allowance) the ultimate saturated average flow is approximately 49,700 m3/d, i.e. about 54,700 m3/d with the 10% STP design margin. Since this exceeds the 20,000 m3/d Phase I threshold defined in the Terms of Reference, phasing of the treatment capacity against realistic build-out horizons (2030 / 2055 / ultimate) becomes the central concept decision of the project.")
bullet("The study area was partitioned into 36 sewer zones, each a contiguous road-network territory draining to a single outlet on the trunk system. Zone boundaries follow the road fabric and plot density rather than pure terrain watersheds.")
bullet("A minimum-gradient depth-accumulation screening (Table 11 gradients, 12 m maximum cover) shows that about 84% of plots reach the STP by gravity. The remaining pockets — mostly wadi-floor villages south-west and north-west of town lying below the STP inlet level — consolidate into 18 candidate sewage lifting stations after cascading nearby pockets together; 147 minor pockets are left to be absorbed by local re-routing during design.")
bullet("The trunk network crosses mapped wadi lines at approximately 134 locations (12 on the main trunk). Crossing design obligations (ductile iron across the crossing plus 15 m either side, 2.0 m minimum cover in soft soil, valve arrangements, MoAFWR approvals and 1:20/1:50/1:100 flood data) are summarised in this report.")
bullet("A structured register of data required from NWS and third parties (population and occupancy statistics, water billing and flow-meter records, existing STP/PS records, as-builts, master plan extracts) is provided; several parameters of this screening are provisional pending that data and shall be confirmed before concept design.")
para("All results are screening-grade: they are derived from an unverified terrain model and cadastral data, and are intended to focus the data collection, survey and modelling effort of the concept stage — not to fix any design parameter.")
pagebreak()

# ================= 2 INTRODUCTION =================
H(1, "Introduction and Project Background")
para("Nama Water Services (NWS/OWWSC) intends to extend wastewater services to all unserved plots within the operational sewer and treated-effluent (TE) service area of Ibri Wilayat, to verify the RG Master Plan 2018, and to develop the concept, preliminary and detailed designs and tender documents for the wastewater network, the TE network and the sewage treatment capacity, for a design horizon equal to the year of project completion plus 25 years, or ultimate (saturated) conditions (TOR Section 03).")
para("The scope explicitly comprises two parallel piped systems: the wastewater collection system (property connections, laterals, main sewers, trunk mains, lifting/pumping stations and rising mains, modelled in SewerGEMS) and the treated-effluent distribution system (TE mains and customer connections from the STP to TE customers, modelled in WaterGEMS), together with the new STP capacity. For the STP, the Terms of Reference make the preliminary design and EPC tender preparation part of this consultancy only where the Phase I design capacity is below 20,000 m3/d; at greater capacities those stages are subject to separate NWS instruction.")
para("This Concept Screening Report is a preparatory engineering assessment. It precedes the formal data collection, topographical survey, geotechnical investigation and hydraulic modelling of the concept stage, and its purpose is threefold: (i) to apply the NWS design criteria to the available data and establish the feasible shape of the collection system; (ii) to identify the decisions and risks that will drive the concept options; and (iii) to define precisely which data must be obtained before the concept design can be fixed.")
H(2, "Design basis documents")
bullet("PAM-GUD-203 — Wastewater Design Guidelines v1.0 (Rev 01). Cited in this report as (p##).")
bullet("PAM-GUD-201 — General Design Guidelines v1.0 (Rev 01). Cited as (G1-p##).")
bullet("Terms of Reference, Section 03, Tender T/2719110/2025 (cited as TOR).")

# ================= 3 SCOPE =================
H(1, "Scope of Services and Deliverables")
para("The consultancy proceeds in stages: Concept Design and associated studies; Preliminary Design; Detailed Design; Tender Documentation and tendering support; followed by construction supervision. The obligations most relevant to the present screening are summarised below.")
H(2, "Key technical obligations")
table(["Obligation", "Requirement", "Reference"], [
 ["Coverage", "All plots within the project boundary — built, open and under construction — shall be provided with sewer and TE connections.", "TOR H.2"],
 ["Gravity preference", "Sewage connection by gravity is preferred; pumping only where gravity is not realistic or is excessively costly, with both options considered.", "TOR p10"],
 ["Topography-led layout", "The network layout shall utilise gravity as far as practicable using contour mapping to avoid pumping.", "TOR p12"],
 ["Options", "Minimum three options each for the wastewater network, the TE network, and each STP.", "TOR p13"],
 ["Design horizon", "Completion + 25 years or ultimate (saturated); model years start / 2030 / 2055 / ultimate at 5-year projection intervals.", "TOR p3, p14"],
 ["Modelling platforms", "SewerGEMS (wastewater) and WaterGEMS (TE), with Excel calculation deliverables.", "TOR p14"],
 ["Inverted siphons", "To be avoided; permitted only where no feasible alternative exists.", "TOR p12"],
 ["STP Phase I threshold", "Preliminary design and EPC tender for the new STP fall within this consultancy only if Phase I capacity is below 20,000 m3/d.", "TOR p3, p5"],
], widths=[1.3, 4.2, 1.0])
H(2, "Surveys and investigations in the consultant scope")
para("The designs shall be based on surveys and investigations executed under this consultancy, not on the information folder, which is provided for information only (TOR H.14):")
bullet("Topographical survey of the whole project area, with ground levels, cover levels, inverts and asset data of the existing sewer and TE systems, delivered to NWS as-built/GIS specifications (TOR H.11, Concept deliverables).")
bullet("Geotechnical investigation at concept and design stages: soil profile and strength for excavation methods and shoring, pipe bedding and structural design of deep trunk sewers, manholes and pump stations, groundwater levels for dewatering assessment, and chemical aggressivity for material selection (TOR H.12; deliverables 20-22).")
bullet("Trial pits — fifty (50) at critical locations proposed by the Consultant and approved by NWS: their purpose is the physical identification of existing underground services (position, depth, size) at road intersections, along major utility corridors and along the expected routes of trunk sewers and rising mains, to secure constructability and clash-free alignments; they also verify groundwater conditions locally (TOR H.13, Utility Surveys clause).")
bullet("CCTV and condition surveys of existing assets where rehabilitation or integration is required (per NF EN 13508-2 / GUD-203 Section 13).")
H(2, "Concept-stage deliverables")
para("The concept design report shall include, inter alia: executive summary; data collection and validation report; as-built and GIS of existing sewer and TE assets; population forecast and flow projections at 5-year intervals; wastewater generation, design flows and flow directions; minimum three network options with hydraulic calculations (SewerGEMS/WaterGEMS); concept designs for lifting/pumping stations wherever required; STP location report and concept design with at least three options per STP; TE and sludge management strategy; tanker discharge and TE filling facilities; emergency lagoons (5-day storage) and excess-TE wadi discharge provisions; life-cycle costs; environmental and social screening; contracting strategy and packaging (TOR 4.1.1-4.1.1.2).")

# ================= 4 CRITERIA =================
H(1, "Design Basis and Criteria")
para("This section presents the criteria that governed the screening, with their engineering rationale. All values are taken from the NWS guidelines without modification; where the guidelines defer a value to other sources, this is stated explicitly and the value is carried as an assumption to be confirmed.")
H(2, "Hydraulic design of gravity sewers")
para("Gravity sewers are designed with the Colebrook-White formula using an operational roughness ks = 1.5 mm for all pipe sizes and materials, or alternatively Manning's formula, computed in NWS-approved software (p24). The roughness value represents an aged pipe in sewage service, including joint and slime effects, rather than the catalogue value of the new pipe.")
para("Two velocity limits bound the design. The minimum self-cleansing velocity of 0.75 m/s at peak flow (0.90 m/s preferred) ensures that sand and organic solids do not settle permanently; where early flows cannot reach it — typically at the heads of the system — the guideline requires a minimum tractive-force (boundary shear) check instead, which recognises that a shallow fast film can transport sediment even at low discharge (p26-27). The maximum velocity of 3.0 m/s protects the pipe invert and structures from abrasion and prevents excessive turbulence, splashing and release of hydrogen sulphide (p27).")
para("At peak flow the proportional depth d/D is limited to 0.65 for pipes up to DN350 and 0.50 above DN350 (p27). The reserve above the flow surface ventilates the sewer, limits septicity, and retains capacity for flows beyond the design estimate; the stricter limit for large sewers reflects the greater consequence of surcharge in trunk mains.")
para("These hydraulic requirements translate into the minimum gradients of Table 11 (p29), reproduced below; the steeper of the Table 11 value and the tractive-force value governs. Sewers shall not be oversized to justify flatter gradients, and uniform gradients are required between manholes (p29).")
table(["Sewer diameter (mm)", "Minimum gradient (mm/m)"], [
 ["200", "5.00"], ["250", "3.75"], ["315", "2.70"], ["400", "2.05"],
 ["500", "1.55"], ["600", "1.25"], ["700", "1.00"], ["800", "0.85"], ["900 and above", "0.75"],
], widths=[3.2, 3.2])
caption("Table — Minimum sewer gradients, PAM-GUD-203 Table 11 (p29)")
H(2, "Depth, cover and manholes")
para("The minimum depth to pipe crown is 1.3 m (0.5 m where concrete protection is provided), protecting the pipe from traffic loads and utility conflicts (p33). The recommended maximum cover is approximately 10-12 m: beyond it, excavation, shoring and dewatering costs escalate and maintenance access deteriorates, and the guideline directs the designer to introduce a pumping station rather than deepen the sewer further (p33). This 12 m ceiling, combined with the Table 11 gradients, is the criterion by which the present screening decides where gravity fails.")
para("Manholes are required at every change of gradient, diameter or alignment, at junctions and lateral ends, at maximum spacings of 100 m (DN200-315), 120 m (DN350-900), 150 m (DN1000-1400) and 200 m above DN1400 (p30). Drops above 600 mm require external backdrops, limited to 2 m height (p30). Manholes and pipelines shall not be located in wadis or areas subject to washout (p30, p33).")
H(2, "Trunk mains")
para("NWS defines trunk mains as gravity pipelines larger than DN800, longer than 1,000 m without connections, immediately upstream of the STP or a main pumping station (p35). Materials above DN600 are GRP, lined reinforced concrete, or profile-wall HDPE (p35). Service corridors range from 2.0 m (up to DN500) to 4.4 m (DN2400) (p32-35).")
H(2, "Pumping stations and force mains")
para("Where pumping is unavoidable, stations are sited on hydraulic merit — flooded suction, minimised energy, force-main economics (p38) — with duty/standby pumps each capable of the design flow in small stations (p39). Force mains maintain 0.75 m/s minimum (1.0 m/s where flow is intermittent), 2.5 m/s maximum; gradients not flatter than 1:750, ideally 1:500 rising / 1:300 falling; retention below 30 minutes where achievable; air valves at high points and washouts at low points; access every 500 m; 3.0 m horizontal clearance from water mains, crossing beneath them with 450 mm vertical separation (p50-51).")
H(2, "Wadi crossings")
para("Crossing design requires hydrological data (wadi profiles, 1:20, 1:50 and 1:100 year floods, bed material, bed-level trends) obtained from CAA and MoAFWR, with MoAFWR approval (G1-p85). Ductile iron pipe with restrained joints is used across the crossing and 15 m beyond each bank; protection follows standard drawing PAM-STD-404; anti-flotation is checked for the empty pipe under flood; minimum cover in soft soil is 2.0 m; isolation and air valves are placed on both banks with a washout at the low point; and no chambers or marker posts are permitted in the wadi bed or embankments (G1-p86). Excess treated effluent discharged to a wadi must meet Class A quality (MD 145/93) and requires Environment Authority approval (p73).")
H(2, "Flow estimation methodology")
para("Flows are computed by the PAM-GUD-201 chain: population is derived from plots (plots x average properties per plot x occupancy rate, with the occupancy rate taken from NCSI housing statistics, G1-p58); domestic water consumption for Adh Dhahirah Governorate is 164 l/c/d (IMP 2024 baseline, G1-p60), increased by the governorate ratios of 22% for non-domestic and 14% for governmental consumption (G1-p60); wastewater generation applies return rates of 85% for domestic and tanker consumption and 54% for non-domestic and governmental consumption (G1-p71); a linear infiltration allowance of 720 L/d per km of new sewer is added (G1-p72); the wastewater peaking factor follows the Peltier formula PfWW = 1.5 + 1/sqrt(Qm) with Qm in L/s, capped at 5.0 (G1-p72); and STP capacity carries an additional 10% operational margin (G1-p73). The design horizon of the networks is completion + 25 years or saturation (TOR), while STPs are designed for phased horizons of at least 15 years (p65).")
para("For the present screening, the occupancy rate has been provisionally taken as 6.0 persons per housing unit with 1 property per plot; both parameters shall be replaced by NCSI figures before concept design, and all flows scale linearly with them.")
pagebreak()

# ================= 5 STUDY AREA =================
H(1, "Study Area")
para("The project boundary covers 439.8 km2 of Ibri Wilayat in Adh Dhahirah Governorate. Urban development concentrates in the central agglomeration of Ibri town, extending along the dual-carriageway corridors north-eastwards (Al Araqi and beyond) and south-eastwards, with satellite settlements to the east and south. The existing Ibri STP occupies a low position at the south-western edge of the town at ground level of approximately 327.5 m.")
table(["Land use class", "Plots", "Area (km2)"], [
 ["Residential", "43,722", "32.2"], ["Agricultural", "4,310", "36.8"], ["Governmental", "2,976", "16.3"],
 ["Commercial", "3,057", "2.4"], ["Industrial", "466", "0.9"], ["Not classified", "3,991", "7.5"],
], widths=[2.6, 1.6, 1.6])
caption("Table — Land use within the project boundary (MoHUP cadastre)")
pic(os.path.join(MAPS, "W2_M1_Study_Area.png"))
caption("Map M1 — Study area, MoH plots and existing STP")
H(2, "Existing sewer and TE system")
para("Figure No. 2 of the Terms of Reference (Existing Sewer and TE Systems) shows that a sewer system is already in operation in the north-eastern district of the agglomeration (Al Araqi area), connected to the existing STP by a trunk corridor running south-west through the town. The concept design must therefore (i) verify the route, capacity, condition and hydraulic performance of this existing trunk and its district network against the design-horizon flows, (ii) integrate or redesign it as required, and (iii) decide whether the new trunk system parallels, reuses or replaces the existing corridor. The as-built records, drawings and previous design documentation of this system are first-priority data items (Section 8). Zones overlapping the served district contain existing customers, whose flows are already reaching the STP and must not be double-counted in phasing.")
para("Terrain falls from the jebel foothills in the north and east towards the south-west. Within the built-up area, ground levels range from about 320 m near the STP to about 490 m at the north-eastern settlements. Wadi channels dissect the entire area and control both the gravity drainage pattern and the crossing inventory of any trunk system. Terrain analysis in this report uses the project digital terrain model (5 m); as the topographic survey is ongoing, all levels are indicative.")
pic(os.path.join(MAPS, "W2_M2_Terrain.png"))
caption("Map M2 — Terrain and drainage (DTM)")
pic(os.path.join(MAPS, "W2_M3_Roads_Streams.png"))
caption("Map M3 — Road network and wadi streams")

# ================= 6 METHODOLOGY =================
H(1, "Screening Methodology")
para("The screening treats the public road network as the universe of feasible sewer corridors, consistent with the guideline requirement that the secondary network follows the street layout (p32). The road centrelines were converted into a connected network of about 13,400 junction nodes and 17,600 street segments. Where a main road is represented by two parallel carriageway centrelines, the pair was collapsed into a single routing corridor, so that a trunk sewer is routed once along the road reserve rather than once per carriageway.")
para("Every serviceable plot was assigned to its nearest street node, giving each node a demand weight in plots. Ground levels at all nodes were sampled from the DTM. From the existing STP, least-cost routes were computed to every node through the street network; the cost function prefers major-road corridors (which offer wide reserves and direct alignments and avoid dense internal streets) and penalises routes that would force sewage to climb, so that the resulting trunk paths track the natural fall of the terrain along main roads with as few directional changes as practicable.")
para("Sewer zones were then formed as road-network territories: each of the 36 selected outlet points on the trunk system claims the street nodes closer to it (through the network) than to any other outlet, weighted by plot density. Zone boundaries therefore follow the road fabric and the actual distribution of plots — not topographic watershed divides — and each zone drains to exactly one outlet, as required for staged development and hydraulic bookkeeping. The zone polygons shown on the maps are the dissolved territories of these nodes, clipped to the built-up envelope.")
para("Gravity feasibility was screened by accumulating pipe invert levels along every route: starting at 1.9 m below ground at the network heads, the invert falls at the Table 11 minimum gradient applicable to the accumulated upstream load (DN200 head sewers at 5.0 mm/m grading down to 0.75 mm/m for trunk-scale flows), riding no shallower than minimum cover where the ground falls faster. Wherever the accumulated depth exceeds the 12 m maximum cover of the guideline (p33), gravity is deemed infeasible for that route and the node is flagged. Contiguous flagged areas were consolidated: each becomes a single candidate sewage lifting station at its lowest point, and nearby stations within 1.5 km were cascaded into one. Pockets below 50 plots were set aside for resolution by local re-routing during design.")
para("The method deliberately over-detects at the margins — it cannot re-grade individual streets or exploit private-land shortcuts — and its outputs are candidates for design, not designed stations. Its value is that every flagged location follows deterministically from the guideline gradients, the cover limit and the terrain, so the concept design can concentrate on eliminating stations rather than finding them.")

# ================= 7 FINDINGS =================
H(1, "Concept Findings")
H(2, "Gravity feasibility and trunk system")
para("The main agglomeration commands a median available fall of approximately 3.6 m/km towards the existing STP, with the main trunk corridor descending monotonically over its 22 km length. Approximately 84% of all plots reach the STP within the guideline gradients and cover limits by gravity alone. The concept trunk system comprises 22 km of main trunk and about 172 km of trunk branches along major-road corridors.")
pic(os.path.join(W2, "img", "trunk_profile.png"), w=6.2)
caption("Figure — Ground long-section along the main trunk corridor (DTM)")
pic(os.path.join(MAPS, "W2_M4_Trunk_Zones.png"))
caption("Map M4 — Concept trunk and sewer zones")
H(2, "Sewer zones and design flows")
para("The 36 zones and their ultimate saturated flows are listed in the Appendix. Aggregate results:")
table(["Quantity", "Value"], [
 ["Serviceable plots (all classes)", "53,503"],
 ["Population at provisional occupancy (6.0 p/hu)", "~279,800 (saturated)"],
 ["Average daily flow Qadf (ultimate)", "~49,700 m3/d"],
 ["Qadf + 10% STP margin", "~54,700 m3/d"],
 ["Peak flow (Peltier, by zone)", "~86,400 m3/d"],
 ["Collection network length (screening)", "~1,290 km streets carrying sewers"],
], widths=[3.8, 2.6])
caption("Table — Ultimate flow summary (provisional occupancy; scales with NCSI data)")
para("The ultimate capacity requirement is therefore roughly 2.5 times the 20,000 m3/d threshold below which the STP preliminary design and EPC tender fall automatically within this consultancy. The immediate consequence is that STP phasing drives the commercial and technical structure of the project: Phase I sizing must follow from the 2030/2055 build-out projections — not from saturation — and those projections require the population, occupancy and development-rate data listed in Section 8. At typical land footprints of 0.9-3.6 m2 per m3/d (GUD-203 Table 28), the ultimate treatment capacity implies 5-20 ha depending on technology, which should be secured at the STP site now even if construction is phased.")
H(2, "Sewage lifting stations")
para("Eighteen candidate SLS remain after consolidation, concentrated in wadi-floor settlements south-west and north-west of the town centre that lie below the STP inlet level or below their surrounding drainage fabric. Together they lift about 16% of the plots. The three largest candidates serve approximately 1,240, 1,070 and 890 plots with screening peak flows of 29, 26 and 19 L/s respectively. During concept design each candidate shall be tested against: (i) elimination by local re-grading or re-routing; (ii) cascading into a neighbouring station; (iii) retention with force main to the nearest gravity trunk; and for the remote eastern settlements, (iv) comparison of long conveyance against local (satellite) treatment on whole-life cost, as the options analysis requires (G1-p95-97).")
pic(os.path.join(MAPS, "W2_M5_SLS.png"))
caption("Map M5 — Proposed sewage lifting stations (candidates)")
H(2, "Wadi crossings")
para("The concept trunk network crosses mapped stream lines at approximately 134 locations, of which 12 lie on the main trunk. Each crossing carries the design obligations summarised in Section 4 (ductile iron across the crossing +15 m, 2.0 m cover in soft soil, valve arrangement, flood analysis and MoAFWR approval). The stream network derives from the terrain model and shall be re-based on the topographic survey; crossings will be a significant cost and approval item and early engagement with MoAFWR is recommended.")
pic(os.path.join(MAPS, "W2_M6_Wadi_Crossings.png"))
caption("Map M6 — Trunk wadi crossings")

# ================= 8 DATA =================
H(1, "Data Required from NWS and Third Parties")
para("The following register lists the data needed to convert this screening into the concept design. Items marked (critical) gate the flow projections and the STP phasing decision.")
table(["#", "Data item", "Source", "Use"], [
 ["1", "Integrated Master Plan 2024 extracts for Ibri: population forecast, per-capita updates, phasing assumptions (critical)", "NWS Planning", "Flow projections 2030/2055"],
 ["2", "NCSI census and housing data at wilayat/settlement level: population, housing units, occupancy rate (critical)", "NCSI via NWS", "Population from plots"],
 ["3", "Water billing / consumption records (3 years) and active account counts by category; electricity account counts for pro-rata distribution", "NWS", "Calibrate demand and occupancy"],
 ["4", "Water network coverage plans and NRW figures for Ibri", "NWS", "Return-rate assumptions"],
 ["5", "Existing sewer network as-builts, GIS, connected customer list, house-connection records", "NWS O&M", "Integration and redesign scope"],
 ["6", "Existing Ibri STP: design capacity, process, headworks invert levels, spare capacity, O&M constraints (critical)", "NWS", "Trunk terminal level; integration"],
 ["7", "STP inflow records — flow meter data, daily/diurnal, minimum 12 months (critical)", "NWS SCADA/O&M", "Calibrate flows and peak factors"],
 ["8", "Influent/effluent quality records (LIMS) for existing STP", "NWS LIMS", "Process design basis"],
 ["9", "Existing lifting/pumping stations: locations, drawings, capacities, SCADA flow data", "NWS", "Reuse/integration assessment"],
 ["10", "Yellow-tanker discharge records at the STP; tanker filling station data", "NWS", "Present-day flow reconciliation"],
 ["11", "TE system: existing network as-builts, current customers and demands, TSE quality records", "NWS", "TE network design"],
 ["12", "RG Master Plan 2018 and Figures F2/F3 boundaries in GIS format", "NWS", "Existing vs new design areas"],
 ["13", "MoHUP: planned land-use changes, subdivision schemes, land bank for STP/SLS sites", "MoHUP", "Zone build-out; site securing"],
 ["14", "Utility records: water, electricity (OETC/MEDC), telecom, gas within the corridor network", "Utilities", "Trial pit targeting; clash analysis"],
 ["15", "Existing geotechnical investigations and groundwater monitoring in Ibri", "NWS / MoT / others", "GI planning"],
 ["16", "Wadi flood studies, aflaj inventory and protection requirements", "MoAFWR / CAA", "Crossing design and approvals"],
 ["17", "Survey control: benchmarks and national datum references in the wilayat", "NSA / NWS", "Topographic survey datum"],
 ["18", "NWS standard drawings (PAM-STD series), as-built/GIS specifications, BD data requirements", "NWS", "Deliverable compliance"],
 ["19", "Existing sewer & TE network files: as-built DWG/CAD and GIS/shapefiles, incl. Figure F2 source data (critical)", "NWS", "Integration; served-area mapping"],
 ["20", "Previous design reports, hydraulic models and calculations of the existing system and trunk to the STP (critical)", "NWS", "Capacity verification"],
 ["21", "Existing trunk main details: diameter, material, invert profile, condition/CCTV records, connection points", "NWS O&M", "Reuse vs replace decision"],
], widths=[0.35, 2.9, 1.2, 2.0], font=8)

# ================= 9 CONCLUSIONS =================
H(1, "Conclusions and Next Steps")
bullet("The gravity concept is confirmed as the backbone strategy: the terrain supports gravity conveyance of the great majority of the catchment to the existing STP site, in line with the TOR preference.")
bullet("STP phasing is the decisive early question; it cannot be settled without the master plan and NCSI data of Section 8, which should be requested at kick-off.")
bullet("The 18 SLS candidates and the remote eastern settlements define the option space for the mandated three network options: maximum-gravity, consolidated-pumping, and satellite-treatment variants.")
bullet("Early engagements recommended: MoAFWR (wadi crossings and flood data), MoHUP (STP/SLS land), NWS O&M (STP records and as-builts).")
bullet("Next steps: mobilise topographic survey and data collection; georeference existing-system extents; construct the SewerGEMS model on the zone structure of this report; develop and cost the three concept options; prepare the STP location and phasing report.")
pagebreak()

# ================= APPENDIX =================
H(1, "Appendix A — Zone Schedule (Ultimate, Saturated)")
zf = list(csv.reader(open(os.path.join(W2, "report", "zone_flows.csv"))))
rows = [[r[0], r[1], r[6], r[7], r[8], r[9], r[10]] for r in zf[1:]]
table(["Zone", "Plots", "Population", "Qadf (m3/d)", "PF", "Qpeak (m3/d)", "Streets (km)"], rows,
      widths=[0.6, 0.8, 1.0, 1.1, 0.6, 1.1, 1.0], font=8)
caption("Population at provisional occupancy 6.0 persons/housing unit; to be updated with NCSI data.")

doc.save(OUT)
print("saved", OUT)
